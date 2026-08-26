"""
MathMate — Interactive Mathematics Lab (Version 2.0)
Structured according to Syllabus Requirements:

UNIT I: Practical based on basics of integers, real numbers and complex numbers
  1. Integers and Divisibility (Prime Factorization, Divisors, Primality)
  2. Computation of greatest common divisor using Euclid’s algorithm
  3. Complex Numbers & Polar Form

UNIT II: Practical based on Introduction to basic counting and basics of functions
  4. Permutations of Distinct Objects (n!, P(n,r), Circular)
  5. Combinations of Distinct Objects (C(n,r), Binomial)
  6. Injective, Bijective, Surjective Functions
  7. Inverse Images of Sets under Functions (f⁻¹(S))
"""

import sys
import os
from pathlib import Path

# Ensure current working directory & file parent are on sys.path for Streamlit Cloud
ROOT_DIR = Path(__file__).parent.resolve()
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

import io
import re
import math
import random
import json
from datetime import datetime, timedelta

import streamlit as st

# ============================================================
# PAGE CONFIG (MUST BE THE FIRST STREAMLIT COMMAND EXECUTED)
# ============================================================
st.set_page_config(
    page_title="MathMate 2.0 — Interactive Mathematics Lab",
    page_icon="🧮",
    layout="wide",
    initial_sidebar_state="expanded"
)

import numpy as np
import plotly.graph_objects as go
import sympy as sp
from sympy import factorint, gcdex, mod_inverse

try:
    import db
    if not hasattr(db, "init_db") or not hasattr(db, "save_solution"):
        raise AttributeError("Imported db is missing required attributes")
except Exception:
    class DummyDB:
        @staticmethod
        def init_db(): pass
        @staticmethod
        def is_supabase_connected(): return False
        @staticmethod
        def save_solution(*args, **kwargs): pass
        @staticmethod
        def fetch_history(*args, **kwargs): return []
        @staticmethod
        def save_user_stats(*args, **kwargs): pass
        @staticmethod
        def fetch_user_stats(): return 0, 0, 0, 0, 0
        @staticmethod
        def load_user_stats(): return {"streak": 0, "xp": 0, "quiz_correct": 0, "quiz_total": 0, "problems_solved": 0, "last_active_date": ""}
    db = DummyDB


# Optional export libs — degrade gracefully
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False


# ============================================================
# STYLES & STYLISH COLOR PALETTE (CSS DESIGN SYSTEM)
# ============================================================
INK = "#14213D"
PAPER = "#F7F5EF"
TEAL = "#2EC4B6"
AMBER = "#FFB627"
CORAL = "#E63946"
GRAPHITE = "#4A4E69"

st.markdown(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=JetBrains+Mono:ital,wght@0,400;0,700;1,400&display=swap');
    
    html, body, [class*="css"] {{
        font-family: 'Outfit', sans-serif;
        background-color: #0A1226;
        color: #F7F5EF;
    }}
    
    .main .block-container {{
        padding-top: 1.5rem;
        padding-bottom: 3rem;
        max-width: 1350px;
    }}
    
    .stApp {{
        background: radial-gradient(circle at 10% 20%, rgba(20, 33, 61, 0.95) 0%, rgba(10, 18, 38, 1) 90%);
    }}
    
    [data-testid="stSidebar"] {{
        background: #0D1B2A !important;
        border-right: 1px solid rgba(255, 255, 255, 0.08);
    }}
    
    .hero-container {{
        text-align: center;
        padding: 2.2rem 1.5rem;
        background: radial-gradient(circle at 50% 30%, rgba(46, 196, 182, 0.15) 0%, rgba(13, 27, 42, 0.65) 80%);
        border-radius: 20px;
        border: 1px solid rgba(46, 196, 182, 0.28);
        margin-bottom: 1.8rem;
        box-shadow: 0 12px 40px rgba(0, 0, 0, 0.35);
    }}
    
    .hero-title {{
        font-size: 2.8rem;
        font-weight: 800;
        background: linear-gradient(135deg, #FFFFFF 0%, #2EC4B6 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0.2rem;
        letter-spacing: -0.02em;
    }}
    
    .hero-subtitle {{
        font-size: 1.05rem;
        color: rgba(247, 245, 239, 0.8);
        font-weight: 400;
        margin-bottom: 1.2rem;
    }}
    
    .glass-card {{
        background: rgba(255, 255, 255, 0.035);
        backdrop-filter: blur(12px);
        -webkit-backdrop-filter: blur(12px);
        border: 1px solid rgba(255, 255, 255, 0.09);
        border-radius: 16px;
        padding: 22px;
        margin-bottom: 18px;
        box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.25);
        transition: transform 0.2s ease, border-color 0.2s ease, box-shadow 0.2s ease;
    }}
    
    .glass-card:hover {{
        border-color: rgba(46, 196, 182, 0.4);
        box-shadow: 0 10px 35px rgba(46, 196, 182, 0.15);
    }}
    
    .topic-card {{
        background: rgba(255, 255, 255, 0.03);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 16px;
        padding: 20px;
        min-height: 230px;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
        transition: all 0.25s cubic-bezier(0.4, 0, 0.2, 1);
    }}
    
    .topic-card:hover {{
        transform: translateY(-4px);
        border-color: rgba(46, 196, 182, 0.45);
        background: rgba(255, 255, 255, 0.05);
        box-shadow: 0 12px 30px rgba(0, 0, 0, 0.4);
    }}

    .topic-icon-lg {{
        font-size: 2.2rem;
        margin-bottom: 6px;
    }}
    
    .unit-badge {{
        display: inline-block;
        padding: 4px 10px;
        border-radius: 20px;
        font-size: 0.72rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.06em;
        margin-bottom: 8px;
    }}

    .unit-badge-1 {{
        background: rgba(46, 196, 182, 0.15);
        color: #2EC4B6;
        border: 1px solid rgba(46, 196, 182, 0.35);
    }}

    .unit-badge-2 {{
        background: rgba(255, 182, 39, 0.15);
        color: #FFB627;
        border: 1px solid rgba(255, 182, 39, 0.35);
    }}

    .topic-meta {{
        font-size: 0.78rem;
        color: rgba(247, 245, 239, 0.55);
        font-weight: 600;
        margin-top: 8px;
        margin-bottom: 12px;
    }}
    
    .timeline-step {{
        display: flex;
        align-items: flex-start;
        margin-bottom: 14px;
        position: relative;
    }}
    
    .timeline-num {{
        width: 32px;
        height: 32px;
        border-radius: 50%;
        background: #2EC4B6;
        color: #0A1226;
        font-weight: 800;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 0.9rem;
        flex-shrink: 0;
        margin-right: 14px;
        box-shadow: 0 0 12px rgba(46, 196, 182, 0.4);
    }}
    
    .timeline-content {{
        background: rgba(255, 255, 255, 0.04);
        border: 1px solid rgba(255, 255, 255, 0.08);
        border-radius: 12px;
        padding: 14px 18px;
        width: 100%;
    }}
    
    .timeline-title {{
        font-weight: 700;
        color: #FFFFFF;
        font-size: 1.05rem;
        margin-bottom: 4px;
    }}
    
    .timeline-body {{
        font-size: 0.92rem;
        color: rgba(247, 245, 239, 0.85);
        line-height: 1.5;
        white-space: pre-wrap;
    }}
    
    .answer-card {{
        background: linear-gradient(135deg, rgba(46, 196, 182, 0.22) 0%, rgba(20, 33, 61, 0.5) 100%);
        border: 2px solid #2EC4B6;
        border-radius: 16px;
        padding: 24px;
        text-align: center;
        margin-bottom: 16px;
        box-shadow: 0 10px 30px rgba(46, 196, 182, 0.25);
    }}
    
    .answer-badge {{
        font-size: 0.8rem;
        font-weight: 800;
        letter-spacing: 0.12em;
        color: #2EC4B6;
        text-transform: uppercase;
        margin-bottom: 6px;
    }}
    
    .answer-value {{
        font-size: 1.85rem;
        font-weight: 800;
        color: #FFFFFF;
        font-family: 'JetBrains Mono', monospace;
    }}
    
    .understand-card {{
        background: rgba(255, 182, 39, 0.06);
        border: 1px solid rgba(255, 182, 39, 0.25);
        border-radius: 14px;
        padding: 18px;
        margin-top: 14px;
    }}
    
    .understand-title {{
        color: #FFB627;
        font-weight: 700;
        font-size: 0.95rem;
        margin-bottom: 8px;
        display: flex;
        align-items: center;
        gap: 6px;
    }}

    .status-pill {{
        display: inline-block;
        padding: 6px 14px;
        border-radius: 8px;
        font-weight: 700;
        font-size: 0.9rem;
        margin: 4px 6px 4px 0;
    }}
    
    .status-success {{
        background: rgba(46, 196, 182, 0.2);
        color: #2EC4B6;
        border: 1px solid #2EC4B6;
    }}

    .status-danger {{
        background: rgba(230, 57, 70, 0.2);
        color: #E63946;
        border: 1px solid #E63946;
    }}
    
    div[data-baseweb="select"] > div {{
        background-color: rgba(255, 255, 255, 0.05) !important;
        border-color: rgba(255, 255, 255, 0.15) !important;
        color: #FFFFFF !important;
        border-radius: 8px !important;
    }}
    
    .stButton > button {{
        border-radius: 10px;
        font-weight: 600;
        transition: all 0.2s ease;
        min-height: 44px;
    }}
    
    .stButton > button[kind="primary"] {{
        background: linear-gradient(135deg, #2EC4B6 0%, #1B9AAA 100%);
        border: none;
        box-shadow: 0 4px 15px rgba(46, 196, 182, 0.3);
    }}
    
    .stButton > button[kind="primary"]:hover {{
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(46, 196, 182, 0.45);
    }}
    
    .hero-symbol-banner {{
        font-size: 2.2rem;
        opacity: 0.18;
        letter-spacing: 0.8rem;
        user-select: none;
        margin-bottom: -6px;
        font-family: 'JetBrains Mono', monospace;
    }}

    .db-status-badge {{
        display: inline-flex;
        align-items: center;
        gap: 6px;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 0.75rem;
        font-weight: 600;
        margin-bottom: 12px;
    }}

    /* Math Formula KaTeX Mobile Scroll & Responsiveness */
    .stKatex, div[data-testid="stKatex"] {{
        overflow-x: auto !important;
        overflow-y: hidden !important;
        max-width: 100% !important;
        padding: 4px 0 !important;
        -webkit-overflow-scrolling: touch;
    }}

    @media (max-width: 768px) {{
        .main .block-container {{
            padding-top: 0.8rem !important;
            padding-bottom: 2.0rem !important;
            padding-left: 0.75rem !important;
            padding-right: 0.75rem !important;
        }}
        
        .hero-title {{
            font-size: 1.95rem !important;
            margin-bottom: 0.3rem !important;
        }}
        
        .hero-subtitle {{
            font-size: 0.9rem !important;
            margin-bottom: 1.0rem !important;
        }}
        
        .hero-symbol-banner {{
            font-size: 1.1rem !important;
            letter-spacing: 0.2rem !important;
            margin-bottom: 0px !important;
        }}
        
        .glass-card {{
            padding: 14px 16px !important;
            margin-bottom: 12px !important;
            border-radius: 12px !important;
        }}

        .topic-card {{
            min-height: auto !important;
            margin-bottom: 12px !important;
        }}
        
        .timeline-step {{
            margin-bottom: 10px !important;
        }}
        
        .timeline-num {{
            width: 26px !important;
            height: 26px !important;
            font-size: 0.8rem !important;
            margin-right: 10px !important;
        }}
        
        .timeline-content {{
            padding: 10px 12px !important;
            border-radius: 10px !important;
        }}
        
        .answer-card {{
            padding: 16px 12px !important;
            margin-top: 12px !important;
            border-radius: 12px !important;
        }}
        
        .answer-value {{
            font-size: 1.35rem !important;
            word-break: break-word !important;
            overflow-x: auto !important;
        }}
        
        .stButton > button {{
            width: 100% !important;
            margin-bottom: 4px !important;
        }}
        
        div[data-testid="stHorizontalBlock"] {{
            gap: 0.5rem !important;
        }}
        
        input, textarea, select {{
            font-size: 16px !important;
        }}
    }}
</style>
""", unsafe_allow_html=True)


# Initialize DB
if hasattr(db, "init_db"):
    try:
        db.init_db()
    except Exception:
        pass


# Initialize & Update User Stats with Correct Date-based Streak
if "streak" not in st.session_state or "xp" not in st.session_state:
    stats = db.load_user_stats() if hasattr(db, "load_user_stats") else {}
    st_streak = stats.get("streak", 0)
    st_xp = stats.get("xp", 0)
    st_quiz_corr = stats.get("quiz_correct", 0)
    st_quiz_tot = stats.get("quiz_total", 0)
    st_probs_solved = stats.get("problems_solved", 0)
    st_last_active = stats.get("last_active_date", "")

    today_str = datetime.now().strftime("%Y-%m-%d")
    yesterday_str = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")

    # Evaluate streak correctly based on consecutive calendar dates
    if st_last_active == today_str:
        calc_streak = st_streak
    elif st_last_active == yesterday_str:
        calc_streak = st_streak
    else:
        calc_streak = 0  # Reset streak if inactive for >1 day

    st.session_state.streak = calc_streak
    st.session_state.xp = st_xp
    st.session_state.quiz_score = {"correct": st_quiz_corr, "total": st_quiz_tot}
    st.session_state.problems_solved = st_probs_solved
    st.session_state.last_active_date = st_last_active
    st.session_state.stats_loaded = True


def record_user_activity(solved_problem=False):
    """Record user activity and maintain correct date-based streak logic."""
    today_str = datetime.now().strftime("%Y-%m-%d")
    yesterday_str = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    last_act = st.session_state.get("last_active_date", "")

    if solved_problem:
        st.session_state.problems_solved = st.session_state.get("problems_solved", 0) + 1

    if last_act != today_str:
        if last_act == yesterday_str:
            st.session_state.streak = st.session_state.get("streak", 0) + 1
        else:
            st.session_state.streak = 1
        st.session_state.last_active_date = today_str

    db.save_user_stats(
        streak=st.session_state.streak,
        xp=st.session_state.xp,
        quiz_correct=st.session_state.quiz_score["correct"],
        quiz_total=st.session_state.quiz_score["total"],
        last_active_date=st.session_state.last_active_date,
        problems_solved=st.session_state.problems_solved
    )


# ============================================================
# TOPIC DEFINITIONS & SYLLABUS METADATA (UNIT I & UNIT II)
# ============================================================
TOPICS = {
    # UNIT I
    "divisibility": "🔢 Integers & Divisibility",
    "gcd": "🧮 Computation of GCD using Euclid’s Algorithm",
    "complex": "📍 Complex Numbers & Polar Form",
    # UNIT II
    "perm": "🔀 Permutations of Distinct Objects",
    "comb": "🎲 Combinations of Distinct Objects",
    "functions": "🔗 Injective, Surjective & Bijective Functions",
    "inverse_image": "🔄 Inverse Images of Sets under Functions",
}

TOPIC_NAV_MAP = {
    "divisibility": "🔢 Integers & Divisibility",
    "gcd": "🧮 Computation of GCD using Euclid’s Algorithm",
    "complex": "📍 Complex Numbers & Polar Form",
    "perm": "🔀 Permutations of Distinct Objects",
    "comb": "🎲 Combinations of Distinct Objects",
    "functions": "🔗 Injective, Surjective & Bijective Functions",
    "inverse_image": "🔄 Inverse Images of Sets under Functions",
}

TOPIC_UNITS = {
    "divisibility": "UNIT I · NUMBER THEORY",
    "gcd": "UNIT I · NUMBER THEORY",
    "complex": "UNIT I · COMPLEX ANALYSIS",
    "perm": "UNIT II · COMBINATORICS",
    "comb": "UNIT II · COMBINATORICS",
    "functions": "UNIT II · FUNCTION THEORY",
    "inverse_image": "UNIT II · FUNCTION THEORY",
}

TOPIC_ICONS = {
    "divisibility": "🔢",
    "gcd": "🧮",
    "complex": "📍",
    "perm": "🔀",
    "comb": "🎲",
    "functions": "🔗",
    "inverse_image": "🔄"
}

TOPIC_KEYWORDS = {
    "divisibility": ["divisibility", "prime factor", "factorization", "divisors", "prime test", "is prime", "factors of"],
    "gcd": [
        "gcd", "hcf", "euclidean", "euclid", "greatest common divisor", "euclid's algorithm", "highest common factor", "bezout",
        "gift pack", "gift packs", "identical pack", "identical packs", "equal group", "equal groups", "items leftover",
        "plank", "planks", "wooden plank", "cut piece", "cut pieces", "without waste", "greatest possible length", "equal smaller pieces",
        "courtyard", "rectangular courtyard", "square tile", "square tiles", "paved entirely", "largest possible size", "tiling",
        "leaving remainder", "leaving remainders", "leaves remainder", "leaves remainders", "remainders of",
        "greatest number that divides", "largest number that divides", "maximum number of"
    ],
    "complex": ["polar form", "modulus", "argument", "rectangular form", "complex number", "argand", "imaginary", "real part"],
    "perm": ["permutation", "permutations", "arrange", "arrangement", "order", "sequence", "line", "row", "npr", "circular permutation"],
    "comb": ["combination", "combinations", "choose", "chosen", "select", "selection", "committee", "team", "pool", "ncr", "ways to choose"],
    "functions": ["injective", "surjective", "bijective", "one-one", "onto", "mapping", "domain", "codomain", "classification"],
    "inverse_image": ["inverse image", "pre-image", "preimage", "f^-1", "f inverse", "inverse of set", "pullback"],
}

TOPIC_CONCEPTS = {
    "divisibility": {
        "title": "Integers & Divisibility",
        "unit": "Unit I",
        "desc": "An integer a divides b (a | b) if b = a·c for some integer c. Every positive integer n > 1 can be uniquely factored into primes.",
        "formula": r"n = p_1^{e_1} \cdot p_2^{e_2} \cdots p_k^{e_k}",
        "identity": r"\tau(n) = \prod (e_i + 1), \quad \sigma(n) = \prod \frac{p_i^{e_i + 1} - 1}{p_i - 1}",
        "key_points": [
            "Fundamental Theorem of Arithmetic guarantees unique prime factorization.",
            "τ(n) gives the total number of positive divisors.",
            "σ(n) gives the sum of all positive divisors."
        ]
    },
    "gcd": {
        "title": "Computation of Greatest Common Divisor using Euclid’s Algorithm",
        "unit": "Unit I",
        "desc": "The Greatest Common Divisor (GCD) of two integers a and b is computed via repeated Euclidean division without prime factorization.",
        "formula": r"a = q \cdot b + r, \quad 0 \le r < b \implies \gcd(a, b) = \gcd(b, r)",
        "identity": r"a \cdot b = \gcd(a, b) \cdot \operatorname{lcm}(a, b) \quad \text{and} \quad a x + b y = \gcd(a, b)",
        "key_points": [
            "Euclidean Algorithm: repeated replacement a = q·b + r until remainder r = 0.",
            "The last non-zero remainder in the division chain is the exact GCD.",
            "Bézout's Identity: Extended Euclidean algorithm finds integers x, y satisfying ax + by = gcd(a,b)."
        ]
    },
    "complex": {
        "title": "Complex Numbers & Polar Form",
        "unit": "Unit I",
        "desc": "Complex numbers z = a + bi can be represented in polar coordinates (r, θ) on the Argand plane.",
        "formula": r"z = r(\cos\theta + i\sin\theta) = r e^{i\theta}",
        "identity": r"r = \sqrt{a^2 + b^2}, \quad \theta = \operatorname{atan2}(b, a)",
        "key_points": [
            "Modulus r is distance from origin on the Argand plane.",
            "Argument θ is angle with positive real axis.",
            "Euler's formula: e^{iθ} = cos θ + i sin θ."
        ]
    },
    "perm": {
        "title": "Permutations of Distinct Objects",
        "unit": "Unit II",
        "desc": "Permutations count ordered arrangements of n distinct objects taken r at a time.",
        "formula": r"P(n, r) = nPr = \frac{n!}{(n-r)!}",
        "identity": r"\text{Full Arrangement: } n!, \quad \text{Circular Permutations: } (n-1)!",
        "key_points": [
            "Order matters for Permutations.",
            "n! counts arrangements of all n distinct items in a row.",
            "Circular arrangements fix 1 position: (n-1)! ways."
        ]
    },
    "comb": {
        "title": "Combinations of Distinct Objects",
        "unit": "Unit II",
        "desc": "Combinations count unordered selections of r distinct objects from n.",
        "formula": r"C(n, r) = nCr = \binom{n}{r} = \frac{n!}{r!(n-r)!}",
        "identity": r"\binom{n}{r} = \binom{n}{n-r}, \quad \sum_{r=0}^n \binom{n}{r} = 2^n",
        "key_points": [
            "Order does NOT matter for Combinations.",
            "Symmetry: C(n, r) = C(n, n-r).",
            "Pascal's Identity: C(n, r) = C(n-1, r-1) + C(n-1, r)."
        ]
    },
    "functions": {
        "title": "Injective, Surjective & Bijective Functions",
        "unit": "Unit II",
        "desc": "Classification of functions based on mapping properties between Domain A and Codomain B.",
        "formula": r"f: A \to B",
        "identity": r"\text{Injective: } f(x_1)=f(x_2) \implies x_1=x_2, \quad \text{Surjective: } \operatorname{range}(f) = B",
        "key_points": [
            "Injective (One-to-One): No two domain elements map to same output.",
            "Surjective (Onto): Every element in codomain has at least one pre-image.",
            "Bijective: Both Injective and Surjective (Invertible)."
        ]
    },
    "inverse_image": {
        "title": "Inverse Images of Sets under Functions",
        "unit": "Unit II",
        "desc": "The inverse image (pre-image) f⁻¹(S) of a subset S ⊆ B is the set of all elements in domain A mapping into S.",
        "formula": r"f^{-1}(S) = \{ x \in A \mid f(x) \in S \}",
        "identity": r"f^{-1}(S_1 \cup S_2) = f^{-1}(S_1) \cup f^{-1}(S_2)",
        "key_points": [
            "f⁻¹(S) lives inside Domain A.",
            "Defined for ANY function, even if non-bijective / non-invertible.",
            "f⁻¹(Ø) = Ø and f⁻¹(Codomain B) = Domain A."
        ]
    }
}


# ============================================================
# HELPER PARSERS & STRING FORMATTERS
# ============================================================
def format_prime_factorization_latex(factor_dict: dict) -> str:
    """Format prime factorization dictionary {p: e} into LaTeX string."""
    if not factor_dict:
        return "1"
    parts = []
    for p in sorted(factor_dict.keys()):
        e = factor_dict[p]
        parts.append(f"{p}^{{{e}}}")
    return r" \times ".join(parts)


def format_prime_factorization_plain(factor_dict: dict) -> str:
    """Format prime factorization dictionary {p: e} into plain readable text."""
    if not factor_dict:
        return "1"
    superscripts = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹'}
    parts = []
    for p in sorted(factor_dict.keys()):
        e = factor_dict[p]
        e_str = "".join(superscripts.get(ch, ch) for ch in str(e))
        parts.append(f"{p}{e_str}")
    return " × ".join(parts)


def clean_math_string(text: str) -> str:
    """Converts raw LaTeX string into clean readable text for HTML cards."""
    superscripts = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵','6':'⁶','7':'⁷','8':'⁸','9':'⁹'}
    t = text.replace(r"\gcd", "gcd").replace(r"\times", "×").replace(r"\cdot", "·").replace(r"\operatorname{lcm}", "lcm").replace(r"\text", "")
    t = re.sub(r"\^\{(\d+)\}", lambda m: "".join(superscripts.get(c, c) for c in m.group(1)), t)
    t = re.sub(r"\^(\d)", lambda m: superscripts.get(m.group(1), m.group(1)), t)
    t = t.replace("{", "").replace("}", "").replace("\\", "").strip()
    return t


def detect_topic(text: str) -> str:
    t = text.lower()
    scores = {k: 0 for k in TOPICS}
    for topic, kws in TOPIC_KEYWORDS.items():
        for kw in kws:
            if kw in t:
                scores[topic] += 1
    best = max(scores, key=scores.get)
    if scores[best] > 0:
        return best
    if any(w in t for w in ["mod", "modulo", "congruence", "≡"]):
        return "congruence"
    elif any(w in t for w in ["factor", "divisible", "prime"]):
        return "divisibility"
    elif any(w in t for w in ["pre-image", "inverse image", "f^-1"]):
        return "inverse_image"
    elif any(w in t for w in ["choose", "select", "committee", "group"]):
        return "comb"
    return "gcd"


def extract_integers(text: str):
    return [int(x) for x in re.findall(r"-?\b\d+\b", text)]


def extract_complex_parts(text: str):
    t = text.replace(" ", "")
    m = re.search(r"z?\=?([+-]?\d+\.?\d*)\s*([+-]\s*\d*\.?\d*)[ij]", t, re.IGNORECASE)
    if m:
        a = float(m.group(1))
        b_str = m.group(2).replace(" ", "")
        b = 1.0 if b_str in ["+", ""] else -1.0 if b_str == "-" else float(b_str)
        return a, b

    m_pure_im = re.search(r"([+-]?\d*\.?\d*)[ij]", t, re.IGNORECASE)
    if m_pure_im and not re.search(r"[+-]\d", t):
        val = m_pure_im.group(1)
        b = 1.0 if val in ["", "+"] else -1.0 if val == "-" else float(val)
        return 0.0, b

    m_real = re.search(r"([+-]?\d+\.?\d*)", t)
    if m_real:
        return float(m_real.group(1)), 0.0
    return 1.0, 1.7320508


def parse_question(text: str):
    topic = detect_topic(text)
    nums = extract_integers(text)

    if topic == "gcd":
        word_kws = ["pack", "plank", "tile", "remainder", "cut", "waste", "courtyard", "piece", "eraser", "pencil", "pen", "divide", "greatest", "maximum", "leftover"]
        is_wp = len(nums) > 2 or any(k in text.lower() for k in word_kws)
        a = nums[0] if len(nums) >= 1 else 1071
        b = nums[1] if len(nums) >= 2 else 462
        return topic, {"a": a, "b": b, "nums": nums, "is_word_problem": is_wp}
    elif topic == "divisibility":
        n = nums[0] if len(nums) >= 1 else 360
        return topic, {"n": abs(n)}
    elif topic == "complex":
        a, b = extract_complex_parts(text)
        return topic, {"a": a, "b": b}
    elif topic in ["perm", "comb"]:
        n = max(nums[0], nums[1]) if len(nums) >= 2 else nums[0] if len(nums) == 1 else 7
        r = min(nums[0], nums[1]) if len(nums) >= 2 else 3
        return topic, {"n": abs(n), "r": abs(r)}
    return topic, {}


# ============================================================
# CORE MATHEMATICAL SOLVER ALGORITHMS
# ============================================================

def solve_divisibility(n: int):
    """UNIT I: Integers and divisibility solver."""
    n = abs(n)
    if n <= 1:
        return [("Input Validation", f"Integer n = {n} must be greater than 1 for prime factorization analysis.")], "n > 1 required", {}

    steps = []
    steps.append(("Input Specification", f"Analyze integer n = {n}"))

    is_prime = sp.isprime(n)
    steps.append(("Primality Test", f"Is {n} prime? {'YES (Prime)' if is_prime else 'NO (Composite)'}"))

    factors = factorint(n)
    fact_plain = format_prime_factorization_plain(factors)
    fact_latex = format_prime_factorization_latex(factors)
    steps.append(("Prime Factorization", f"Factorization: {n} = {fact_plain}"))

    divisors = sorted(sp.divisors(n))
    tau = len(divisors)
    sigma = sum(divisors)

    steps.append(("Divisor Enumeration & Counting Functions",
                  f"Divisors of {n}: {divisors}\n"
                  f"Total Number of Divisors τ({n}) = {tau}\n"
                  f"Sum of Divisors σ({n}) = {sigma}"))

    ans_str = f"{n} = {fact_plain} | τ({n}) = {tau} | σ({n}) = {sigma}"
    ans_latex = fr"{n} = {fact_latex} \implies \tau({n}) = {tau}, \ \sigma({n}) = {sigma}"
    return steps, ans_str, {
        "n": n, "is_prime": is_prime, "factors": factors,
        "divisors": divisors, "tau": tau, "sigma": sigma,
        "fact_plain": fact_plain, "fact_latex": fact_latex,
        "latex_ans": ans_latex
    }


def compute_extended_euclidean_table(a: int, b: int):
    """Computes rows for Extended Euclidean Algorithm Table."""
    a_orig, b_orig = a, b
    rows = []

    r0, r1 = a, b
    x0, x1 = 1, 0
    y0, y1 = 0, 1

    rows.append({
        "Step (i)": 0, "Dividend": a_orig, "Divisor": b_orig,
        "Quotient (q)": "—", "Remainder (r)": r0,
        "x": x0, "y": y0, "is_gcd": False
    })
    rows.append({
        "Step (i)": 1, "Dividend": a_orig, "Divisor": b_orig,
        "Quotient (q)": "—", "Remainder (r)": r1,
        "x": x1, "y": y1, "is_gcd": False
    })

    step_i = 2
    last_gcd_idx = -1
    while r1 != 0:
        q = r0 // r1
        r2 = r0 % r1
        x2 = x0 - q * x1
        y2 = y0 - q * y1

        is_last_nonzero = (r2 == 0)
        rows.append({
            "Step (i)": step_i, "Dividend": r0, "Divisor": r1,
            "Quotient (q)": q, "Remainder (r)": r2,
            "x": x2, "y": y2, "is_gcd": False
        })
        if is_last_nonzero:
            last_gcd_idx = len(rows) - 2

        r0, r1 = r1, r2
        x0, x1 = x1, x2
        y0, y1 = y1, y2
        step_i += 1

    if last_gcd_idx >= 0:
        rows[last_gcd_idx]["is_gcd"] = True

    return rows


def solve_gcd_euclidean(a: int, b: int):
    """UNIT I: GCD using Euclid's algorithm and Extended Bézout identity."""
    a_abs, b_abs = abs(a), abs(b)
    if a_abs == 0 and b_abs == 0:
        return [("Error", "gcd(0, 0) is undefined.")], "Undefined", {}

    steps = []
    if a_abs < b_abs:
        a_abs, b_abs = b_abs, a_abs
        steps.append(("Swap Convention", f"Set a = {a_abs}, b = {b_abs} so that a ≥ b."))

    steps.append(("Initial Values", f"Compute gcd({a_abs}, {b_abs}) using Euclidean division."))

    r_prev, r_curr = a_abs, b_abs
    step_num = 1
    gcd_val = b_abs

    while r_curr != 0:
        q = r_prev // r_curr
        r_next = r_prev % r_curr
        steps.append((f"Division Step {step_num}", f"{r_prev} = {r_curr} × {q} + {r_next}"))
        if r_next != 0:
            gcd_val = r_next
        r_prev, r_curr = r_curr, r_next
        step_num += 1

    if b_abs == 0:
        gcd_val = a_abs

    g, x, y = gcdex(a_abs, b_abs)
    bezout_str = f"{a_abs}({int(x)}) + {b_abs}({int(y)}) = {int(g)}"
    steps.append(("Bézout Identity (Extended Euclid)", f"Extended coefficients: {bezout_str}"))

    table_rows = compute_extended_euclidean_table(a_abs, b_abs)
    lcm_val = (a_abs * b_abs) // gcd_val if gcd_val else 0

    ans_str = f"gcd({a_abs}, {b_abs}) = {gcd_val} | Bézout: {bezout_str}"
    ans_latex = fr"\gcd({a_abs}, {b_abs}) = {gcd_val} \implies {a_abs}({int(x)}) + {b_abs}({int(y)}) = {gcd_val}"
    return steps, ans_str, {
        "gcd": int(gcd_val), "x": int(x), "y": int(y),
        "lcm": int(lcm_val), "table_rows": table_rows,
        "latex_ans": ans_latex, "a": a_abs, "b": b_abs
    }


def solve_gcd_multi(nums: list):
    """Computes GCD of multiple integers gcd(a, b, c, ...)."""
    clean_nums = [abs(x) for x in nums if x != 0]
    if not clean_nums:
        return [("Error", "No non-zero integers provided.")], "Undefined", {}
    if len(clean_nums) == 1:
        return solve_divisibility(clean_nums[0])

    steps = []
    steps.append(("Input Specification", f"Compute GCD for set of {len(clean_nums)} numbers: {clean_nums}"))

    current_gcd = clean_nums[0]
    for i in range(1, len(clean_nums)):
        nxt = clean_nums[i]
        sub_steps, sub_ans, sub_extra = solve_gcd_euclidean(current_gcd, nxt)
        new_gcd = sub_extra.get("gcd", math.gcd(current_gcd, nxt))
        steps.append((f"Pairwise Step {i}: gcd({current_gcd}, {nxt})", f"Result = {new_gcd}"))
        current_gcd = new_gcd

    ans_str = f"gcd({', '.join(map(str, clean_nums))}) = {current_gcd}"
    ans_latex = fr"\gcd({', '.join(map(str, clean_nums))}) = {current_gcd}"
    return steps, ans_str, {"gcd": current_gcd, "nums": clean_nums, "latex_ans": ans_latex}


def solve_gcd_word_problem(text: str):
    """Parses real-world word problems and resolves via GCD."""
    nums = extract_integers(text)
    clean_nums = [abs(x) for x in nums if x > 0]
    if len(clean_nums) < 2:
        clean_nums = [1071, 462]

    steps = []
    steps.append(("Problem Formulation", f"Extracted key dimensions/counts: {clean_nums}"))
    steps.append(("Mathematical Modeling", "This word problem requires computing the Greatest Common Divisor (GCD) to find the largest equal grouping / length without leftover."))

    sub_steps, sub_ans, sub_extra = solve_gcd_multi(clean_nums)
    steps.extend(sub_steps[1:])
    g_val = sub_extra.get("gcd", 1)

    steps.append(("Practical Conclusion", f"The maximum size / number of identical groups that can be formed is exactly {g_val}."))
    ans_str = f"Maximum Equal Size / Groups = {g_val}"
    ans_latex = fr"\text{{Optimal Grouping Size}} = \gcd({', '.join(map(str, clean_nums))}) = {g_val}"
    return steps, ans_str, {"gcd": g_val, "nums": clean_nums, "latex_ans": ans_latex}


def solve_complex_to_polar(a: float, b: float):
    """UNIT I: Complex numbers & Argand plane conversion."""
    steps = []
    steps.append(("Complex Number Input", f"z = {a} + {b}i (Rectangular Form)"))

    r = math.hypot(a, b)
    steps.append(("Modulus Calculation", f"r = |z| = √(a² + b²) = √(({a})² + ({b})²) = {r:.4f}"))

    theta_rad = math.atan2(b, a)
    theta_deg = math.degrees(theta_rad)
    steps.append(("Argument Calculation", f"θ = atan2({b}, {a}) = {theta_rad:.4f} rad = {theta_deg:.2f}°"))

    cos_val = math.cos(theta_rad)
    sin_val = math.sin(theta_rad)
    polar_str = f"{r:.4f} (cos {theta_deg:.2f}° + i sin {theta_deg:.2f}°)"
    exp_str = f"{r:.4f} · e^({theta_deg:.2f}° i)"

    steps.append(("Polar Form Assembly", f"z = {polar_str}"))

    ans_str = f"z = {r:.4f} ∠ {theta_deg:.2f}°"
    ans_latex = fr"z = {r:.4f} \left(\cos({theta_deg:.2f}^\circ) + i\sin({theta_deg:.2f}^\circ)\right) = {r:.4f} e^{{{theta_deg:.2f}^\circ i}}"
    return steps, ans_str, {
        "a": a, "b": b, "r": r, "theta_rad": theta_rad,
        "theta_deg": theta_deg, "polar_str": polar_str,
        "exp_str": exp_str, "latex_ans": ans_latex
    }


def solve_perm(n: int, r: int):
    """UNIT II: Permutations of distinct objects."""
    if r < 0 or n < 0 or r > n:
        return [("Validation Error", f"Invalid parameters: n={n}, r={r}. Requirements: 0 ≤ r ≤ n.")], "Invalid", {}

    steps = []
    steps.append(("Permutation Formula", f"P({n}, {r}) = n! / (n - r)!"))

    n_fact = math.factorial(n)
    nr_fact = math.factorial(n - r)
    npr = n_fact // nr_fact

    steps.append(("Factorial Computations", f"{n}! = {n_fact}\n({n} - {r})! = {n - r}! = {nr_fact}"))
    steps.append(("Exact Calculation", f"P({n}, {r}) = {n_fact} / {nr_fact} = {npr}"))

    circ_perm = math.factorial(n - 1) if n >= 1 else 1
    steps.append(("Circular Permutations Note", f"If arranging all {n} distinct objects in a circle: (n - 1)! = ({n}-1)! = {circ_perm}"))

    ans_str = f"P({n}, {r}) = {npr} | Circular({n}) = {circ_perm}"
    ans_latex = fr"P({n}, {r}) = \frac{{{n}!}}{{({n}-{r})!}} = {npr}"
    return steps, ans_str, {
        "n": n, "r": r, "npr": npr, "circ_perm": circ_perm,
        "latex_ans": ans_latex
    }


def solve_comb(n: int, r: int):
    """UNIT II: Combinations of distinct objects."""
    if r < 0 or n < 0 or r > n:
        return [("Validation Error", f"Invalid parameters: n={n}, r={r}. Requirements: 0 ≤ r ≤ n.")], "Invalid", {}

    steps = []
    steps.append(("Combination Formula", f"C({n}, {r}) = n! / (r! × (n - r)!)"))

    n_fact = math.factorial(n)
    r_fact = math.factorial(r)
    nr_fact = math.factorial(n - r)
    ncr = math.comb(n, r)

    steps.append(("Factorial Computations", f"{n}! = {n_fact}\n{r}! = {r_fact}\n({n} - {r})! = {nr_fact}"))
    steps.append(("Exact Calculation", f"C({n}, {r}) = {n_fact} / ({r_fact} × {nr_fact}) = {ncr}"))
    steps.append(("Combinatorial Symmetry Property", f"C({n}, {r}) = C({n}, {n-r}) = {ncr}"))

    ans_str = f"C({n}, {r}) = {ncr}"
    ans_latex = fr"C({n}, {r}) = \binom{{{n}}}{{{r}}} = \frac{{{n}!}}{{{r}!({n}-{r})!}} = {ncr}"
    return steps, ans_str, {"n": n, "r": r, "ncr": ncr, "latex_ans": ans_latex}


def solve_functions(domain: list, codomain: list, mapping: dict):
    """UNIT II: Function classification (Injective, Surjective, Bijective)."""
    steps = []
    steps.append(("Function Specification", f"Domain A = {{{', '.join(map(str, domain))}}}\nCodomain B = {{{', '.join(map(str, codomain))}}}"))

    map_formatted = ", ".join(f"f({k}) = {v}" for k, v in mapping.items())
    steps.append(("Defined Mapping", map_formatted))

    # Injective check
    mapped_vals = list(mapping.values())
    is_injective = (len(mapped_vals) == len(set(mapped_vals)))
    inj_reason = "All domain elements map to distinct codomain outputs." if is_injective else f"Duplicate mapping detected: multiple inputs map to same output."
    steps.append(("Injectivity (One-to-One) Test", f"Injective: {'YES' if is_injective else 'NO'}\nReason: {inj_reason}"))

    # Surjective check
    range_set = set(mapped_vals)
    codomain_set = set(codomain)
    is_surjective = (range_set == codomain_set)
    surj_reason = "Range equals Codomain (every codomain element has a pre-image)." if is_surjective else f"Unmapped codomain elements: {list(codomain_set - range_set)}"
    steps.append(("Surjectivity (Onto) Test", f"Surjective: {'YES' if is_surjective else 'NO'}\nReason: {surj_reason}"))

    # Bijective check
    is_bijective = (is_injective and is_surjective)
    steps.append(("Bijectivity Test", f"Bijective: {'YES' if is_bijective else 'NO'}\nFunction is {'Invertible' if is_bijective else 'Not Invertible'}."))

    class_type = "Bijective (One-to-One & Onto)" if is_bijective else "Injective (One-to-One)" if is_injective else "Surjective (Onto)" if is_surjective else "Neither Injective nor Surjective"
    ans_str = f"Function Type: {class_type}"
    ans_latex = fr"\text{{Function Type: }} \mathbf{{{class_type}}}"
    return steps, ans_str, {
        "is_injective": is_injective, "is_surjective": is_surjective,
        "is_bijective": is_bijective, "range_set": list(range_set),
        "unmapped": list(codomain_set - range_set), "latex_ans": ans_latex
    }


def solve_inverse_image(domain: list, codomain: list, mapping: dict, target_set: list):
    """UNIT II: Inverse images of sets under functions f⁻¹(S)."""
    steps = []
    target_set_clean = [str(x).strip() for x in target_set]
    steps.append(("Set Input", f"Target Subset S ⊆ B: {{{', '.join(target_set_clean)}}}"))

    inv_image = [str(k) for k, v in mapping.items() if str(v) in target_set_clean]
    steps.append(("Pre-Image Definition", r"f⁻¹(S) = { x ∈ Domain A | f(x) ∈ Subset S }"))

    eval_details = []
    for k, v in mapping.items():
        in_s = str(v) in target_set_clean
        eval_details.append(f"f({k}) = {v} {'∈ S' if in_s else '∉ S'}")
    steps.append(("Element-by-Element Evaluation", "\n".join(eval_details)))

    inv_str = f"{{{', '.join(inv_image)}}}" if inv_image else "Ø (Empty Set)"
    steps.append(("Inverse Image Result", f"f⁻¹({{{', '.join(target_set_clean)}}}) = {inv_str}"))

    ans_str = f"f⁻¹(S) = {inv_str}"
    ans_latex = fr"f^{{-1}}(\{{{', '.join(target_set_clean)}\}}) = \{{\{', '.join(inv_image)}\}}" if inv_image else fr"f^{{-1}}(\{{{', '.join(target_set_clean)}\}}) = \emptyset"
    return steps, ans_str, {
        "target_set": target_set_clean,
        "inverse_image": inv_image,
        "latex_ans": ans_latex
    }


def generate_procedural_question():
    """Generates procedural math practice questions across Unit I and Unit II."""
    topics_list = list(TOPICS.keys())
    t = random.choice(topics_list)

    if t == "gcd":
        a = random.randint(12, 120) * random.randint(2, 6)
        b = random.randint(12, 120) * random.randint(2, 6)
        ans = math.gcd(a, b)
        options = sorted(list({ans, ans + random.randint(1, 4), max(1, ans - random.randint(1, 4)), ans * 2}))
        return {
            "topic": t,
            "q": f"Compute the Greatest Common Divisor gcd({a}, {b}) using Euclid's division algorithm.",
            "options": [str(x) for x in options],
            "answer": str(ans),
            "exp": f"Applying Euclid's algorithm: gcd({a}, {b}) = {ans}."
        }
    elif t == "divisibility":
        n = random.choice([24, 36, 48, 60, 72, 90, 100, 120, 180, 360])
        tau = len(sp.divisors(n))
        options = sorted(list({tau, tau + 1, max(1, tau - 2), tau + 3}))
        return {
            "topic": t,
            "q": f"How many positive integer divisors τ(n) does n = {n} have?",
            "options": [str(x) for x in options],
            "answer": str(tau),
            "exp": f"Prime factorization of {n} yields factor powers. τ({n}) = {tau}."
        }
    elif t == "complex":
        a, b = random.choice([(3, 4), (1, 1.732), (5, 12), (1, 1), (0, 4)])
        r = round(math.hypot(a, b), 2)
        options = sorted(list({r, round(r + 1.5, 2), round(max(0.5, r - 1.2), 2), round(r * 1.5, 2)}))
        return {
            "topic": t,
            "q": f"Find the modulus r = |z| for complex number z = {a} + {b}i.",
            "options": [str(x) for x in options],
            "answer": str(r),
            "exp": f"Modulus r = √(a² + b²) = √({a}² + {b}²) = {r}."
        }
    elif t == "perm":
        n = random.randint(5, 8)
        r = random.randint(2, 4)
        ans = math.perm(n, r)
        options = sorted(list({ans, ans + 10, max(1, ans - 12), ans * 2}))
        return {
            "topic": t,
            "q": f"Calculate the number of permutations P({n}, {r}).",
            "options": [str(x) for x in options],
            "answer": str(ans),
            "exp": f"P({n}, {r}) = {n}! / ({n}-{r})! = {ans}."
        }
    elif t == "comb":
        n = random.randint(5, 9)
        r = random.randint(2, 4)
        ans = math.comb(n, r)
        options = sorted(list({ans, ans + 4, max(1, ans - 5), ans + 10}))
        return {
            "topic": t,
            "q": f"Calculate the number of combinations C({n}, {r}).",
            "options": [str(x) for x in options],
            "answer": str(ans),
            "exp": f"C({n}, {r}) = {n}! / ({r}! × ({n}-{r})!) = {ans}."
        }
    else:
        return {
            "topic": "functions",
            "q": "If Domain A has 3 elements and Codomain B has 3 elements, and f maps each domain element to a distinct codomain element, what is f?",
            "options": ["Bijective", "Surjective only", "Injective only", "Neither"],
            "answer": "Bijective",
            "exp": "Since a finite map between equal sized sets is one-to-one, it is also onto, hence Bijective."
        }


# ============================================================
# PLOTLY INTERACTIVE VISUALIZATIONS
# ============================================================
def plot_function_diagram_plotly(domain: list, codomain: list, mapping: dict, highlight_target_set: list = None):
    """Plotly bipartite function mapping diagram."""
    fig = go.Figure()

    d_x, d_y = [0] * len(domain), list(range(len(domain), 0, -1))
    c_x, c_y = [1] * len(codomain), list(range(len(codomain), 0, -1))

    domain_pos = {d: (0, y) for d, y in zip(domain, d_y)}
    codomain_pos = {c: (1, y) for c, y in zip(codomain, c_y)}

    for d, c in mapping.items():
        if d in domain_pos and c in codomain_pos:
            x0, y0 = domain_pos[d]
            x1, y1 = codomain_pos[c]
            is_hl = highlight_target_set and str(c) in highlight_target_set
            line_color = "#FFB627" if is_hl else "#2EC4B6"
            fig.add_trace(go.Scatter(
                x=[x0, x1], y=[y0, y1],
                mode='lines+markers',
                line=dict(color=line_color, width=3 if is_hl else 2),
                hoverinfo='text',
                text=f"f({d}) = {c}",
                showlegend=False
            ))

    fig.add_trace(go.Scatter(
        x=[0] * len(domain), y=d_y,
        mode='markers+text',
        marker=dict(size=28, color='#14213D', line=dict(color='#2EC4B6', width=2)),
        text=domain, textposition="left center",
        textfont=dict(color='#FFFFFF', size=14, family='Outfit'),
        name="Domain A", showlegend=False
    ))

    fig.add_trace(go.Scatter(
        x=[1] * len(codomain), y=c_y,
        mode='markers+text',
        marker=dict(size=28, color='#14213D', line=dict(color='#FFB627', width=2)),
        text=codomain, textposition="right center",
        textfont=dict(color='#FFFFFF', size=14, family='Outfit'),
        name="Codomain B", showlegend=False
    ))

    fig.update_layout(
        title="Bipartite Function Mapping Diagram f: A → B",
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False, range=[-0.4, 1.4]),
        yaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        height=320,
        margin=dict(l=20, r=20, t=40, b=20)
    )
    return fig


def plot_argand_diagram_plotly(a: float, b: float):
    """Plotly Argand plane diagram for complex numbers."""
    r = math.hypot(a, b)
    theta_deg = math.degrees(math.atan2(b, a))

    max_val = max(abs(a), abs(b), 1.0) * 1.3

    fig = go.Figure()

    # Real and Imaginary Axes
    fig.add_trace(go.Scatter(x=[-max_val, max_val], y=[0, 0], mode='lines', line=dict(color='rgba(255,255,255,0.2)', width=1.5), showlegend=False))
    fig.add_trace(go.Scatter(x=[0, 0], y=[-max_val, max_val], mode='lines', line=dict(color='rgba(255,255,255,0.2)', width=1.5), showlegend=False))

    # Vector line from origin
    fig.add_trace(go.Scatter(x=[0, a], y=[0, b], mode='lines+markers', line=dict(color='#2EC4B6', width=3), marker=dict(size=[0, 10], color='#2EC4B6'), name="z = a + bi"))

    # Dotted projection lines
    fig.add_trace(go.Scatter(x=[a, a], y=[0, b], mode='lines', line=dict(color='#FFB627', width=1.5, dash='dash'), showlegend=False))
    fig.add_trace(go.Scatter(x=[0, a], y=[b, b], mode='lines', line=dict(color='#FFB627', width=1.5, dash='dash'), showlegend=False))

    fig.update_layout(
        title=f"Argand Plane: z = {a} + {b}i (Modulus r = {r:.2f}, θ = {theta_deg:.1f}°)",
        xaxis=dict(title="Real Axis (Re)", gridcolor='rgba(255,255,255,0.05)', range=[-max_val, max_val]),
        yaxis=dict(title="Imaginary Axis (Im)", gridcolor='rgba(255,255,255,0.05)', range=[-max_val, max_val]),
        paper_bgcolor='rgba(0,0,0,0)',
        plot_bgcolor='rgba(0,0,0,0)',
        font=dict(color='#F7F5EF', family='Outfit'),
        height=380,
        margin=dict(l=20, r=20, t=50, b=30)
    )
    return fig


# ============================================================
# UI RENDER COMPONENTS
# ============================================================
def render_step_timeline(steps: list):
    """Renders step-by-step reasoning timeline."""
    for i, step in enumerate(steps, 1):
        if isinstance(step, (tuple, list)) and len(step) >= 2:
            title, body = step[0], step[1]
        else:
            title, body = f'Step {i}', str(step)

        st.markdown(f"""
        <div class="timeline-step">
            <div class="timeline-num">{i}</div>
            <div class="timeline-content">
                <div class="timeline-title">{title}</div>
                <div class="timeline-body">{body}</div>
            </div>
        </div>
        """, unsafe_allow_html=True)


def render_euclidean_html_table(table_rows: list):
    """Renders styled HTML table for Extended Euclidean division steps."""
    html = """
    <div style="overflow-x: auto; margin-top: 10px; margin-bottom: 15px;">
        <table style="width: 100%; border-collapse: collapse; background: rgba(255,255,255,0.02); border-radius: 10px; font-size: 0.9rem;">
            <thead>
                <tr style="border-bottom: 2px solid rgba(46,196,182,0.4); text-align: center; color: #2EC4B6;">
                    <th style="padding: 10px;">Step (i)</th>
                    <th style="padding: 10px;">Dividend</th>
                    <th style="padding: 10px;">Divisor</th>
                    <th style="padding: 10px;">Quotient (q)</th>
                    <th style="padding: 10px;">Remainder (r)</th>
                    <th style="padding: 10px;">x</th>
                    <th style="padding: 10px;">y</th>
                </tr>
            </thead>
            <tbody>
    """
    for r in table_rows:
        bg_style = "background: rgba(46,196,182,0.15); font-weight: 700; color: #FFFFFF;" if r.get("is_gcd") else ""
        html += f"""
        <tr style="border-bottom: 1px solid rgba(255,255,255,0.06); text-align: center; {bg_style}">
            <td style="padding: 8px;">{r.get('Step (i)')}</td>
            <td style="padding: 8px;">{r.get('Dividend')}</td>
            <td style="padding: 8px;">{r.get('Divisor')}</td>
            <td style="padding: 8px;">{r.get('Quotient (q)')}</td>
            <td style="padding: 8px; color: {'#2EC4B6' if r.get('is_gcd') else '#FFFFFF'}; font-weight: 700;">{r.get('Remainder (r)')} {'(GCD)' if r.get('is_gcd') else ''}</td>
            <td style="padding: 8px;">{r.get('x')}</td>
            <td style="padding: 8px;">{r.get('y')}</td>
        </tr>
        """
    html += "</tbody></table></div>"
    st.markdown(html, unsafe_allow_html=True)


def render_understand_panel(topic_key: str):
    """Renders theory understand panel card for syllabus topic."""
    c = TOPIC_CONCEPTS.get(topic_key)
    if not c:
        return
    st.markdown(f"""
    <div class="understand-card">
        <div class="understand-title">💡 Understand the Concept: {c['title']}</div>
        <div style="font-size:0.88rem; color:rgba(247,245,239,0.85); line-height:1.5; margin-bottom:8px;">{c['desc']}</div>
    </div>
    """, unsafe_allow_html=True)
    with st.expander("📌 View Core Formula & Key Properties", expanded=False):
        st.markdown("**Core Formula:**")
        st.latex(c['formula'])
        st.markdown("**Key Identity / Relationship:**")
        st.latex(c['identity'])
        st.markdown("**Key Takeaways:**")
        for pt in c['key_points']:
            st.markdown(f"- {pt}")


def format_full_solution_text(question_text: str, topic_name: str, steps: list, answer: str) -> str:
    lines = [
        f"MathMate — Interactive Mathematics Lab",
        f"Topic: {topic_name}",
        f"Question: {question_text}",
        "=" * 50,
        "STEP-BY-STEP REASONING:"
    ]
    for i, s in enumerate(steps, 1):
        if isinstance(s, (tuple, list)) and len(s) >= 2:
            lines.append(f"\nStep {i}: {s[0]}\n{s[1]}")
        else:
            lines.append(f"\nStep {i}: {s}")
    lines.append("=" * 50)
    lines.append(f"FINAL ANSWER: {answer}")
    return "\n".join(lines)


def build_docx(question_text: str, topic_name: str, steps: list, answer: str):
    if not DOCX_AVAILABLE:
        return None
    doc = Document()
    doc.add_heading("MathMate — Interactive Mathematics Lab", level=0)
    doc.add_heading(f"Topic: {topic_name}", level=2)
    doc.add_paragraph(f"Question Statement:\n{question_text}")
    doc.add_heading("Step-by-Step Reasoning:", level=2)
    if steps:
        for i, step in enumerate(steps, 1):
            if isinstance(step, (tuple, list)) and len(step) >= 2:
                title, body = step[0], step[1]
            else:
                title, body = f'Step {i}', str(step)
            doc.add_paragraph(f"Step {i}: {title}", style='Heading 3')
            doc.add_paragraph(str(body))
    doc.add_heading(f"Final Answer: {answer}", level=2)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf(question_text: str, topic_name: str, steps: list, answer: str):
    if not REPORTLAB_AVAILABLE:
        return None
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    subtitle_style = styles['Heading2']
    body_style = styles['Normal']

    story = []
    story.append(Paragraph("<b>MathMate — Interactive Mathematics Lab</b>", title_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"<b>Topic:</b> {topic_name}", subtitle_style))
    story.append(Spacer(1, 10))
    display_q = re.sub(r'<[^>]+>', '', str(question_text)).strip() if question_text else f"{topic_name} Practical Question"
    story.append(Paragraph(f"<b>Question Statement:</b><br/>{display_q}", body_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph("<b>Step-by-Step Reasoning:</b>", subtitle_style))
    story.append(Spacer(1, 8))
    if steps:
        for i, step in enumerate(steps, 1):
            if isinstance(step, (tuple, list)) and len(step) >= 2:
                title, body = step[0], step[1]
            else:
                title, body = f'Step {i}', str(step)
            clean_title = re.sub(r'<[^>]+>', '', str(title))
            clean_body = re.sub(r'<[^>]+>', '', str(body)).replace('\n', '<br/>')
            story.append(Paragraph(f"<b>Step {i}: {clean_title}</b>", body_style))
            story.append(Paragraph(f"{clean_body}", body_style))
            story.append(Spacer(1, 6))

    clean_ans = re.sub(r'<[^>]+>', '', str(answer))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"<b>Final Answer: {clean_ans}</b>", subtitle_style))

    doc.build(story)
    buf.seek(0)
    return buf


def render_answer_card(answer: str, latex_ans: str = None, question_text: str = "", topic_name: str = "", steps: list = None):
    display_ans = clean_math_string(answer)
    st.markdown(f"""
    <div class="answer-card">
        <div class="answer-badge">🎯 FINAL ANSWER</div>
        <div class="answer-value">{display_ans}</div>
        <div style="font-size:0.82rem; color:#2EC4B6; font-weight:700; margin-top:6px;">✓ Computed & Verified</div>
    </div>
    """, unsafe_allow_html=True)
    if latex_ans:
        st.latex(latex_ans)


def set_nav(target_section, target_topic=None, preset_q=None):
    st.session_state.main_nav = target_section
    if target_topic:
        st.session_state.selected_topic_key = target_topic
    if preset_q:
        st.session_state.preset_question = preset_q


# ============================================================
# SIDEBAR NAVIGATION (NESTED HIERARCHICAL RESTRUCTURE)
# ============================================================
st.sidebar.markdown("""
<div style="text-align: left; padding: 4px 0 12px 0;">
    <div style="font-size: 1.6rem; font-weight: 800; color: #2EC4B6; letter-spacing: -0.02em;">🧮 MATHMATE</div>
    <div style="font-size: 0.78rem; color: rgba(247,245,239,0.65);">Interactive Mathematics Lab</div>
</div>
""", unsafe_allow_html=True)

if db.is_supabase_connected():
    st.sidebar.markdown('<div class="db-status-badge" style="background:rgba(46,196,182,0.2); color:#2EC4B6; border:1px solid #2EC4B6;">🟢 Supabase Cloud Active</div>', unsafe_allow_html=True)
else:
    st.sidebar.markdown('<div class="db-status-badge" style="background:rgba(255,182,39,0.2); color:#FFB627; border:1px solid #FFB627;">🟡 SQLite Local Fallback</div>', unsafe_allow_html=True)

if "main_nav" not in st.session_state:
    st.session_state.main_nav = "🏠 Home"

main_nav_options = [
    "🏠 Home",
    "📚 LEARN",
    "🧠 Practice Arena",
    "📐 Formula Bank",
    "📜 Solution History"
]

selected_main = st.sidebar.radio("Navigation", main_nav_options, key="main_nav", label_visibility="collapsed")

selected_topic_key = st.session_state.get("selected_topic_key", "gcd")

if selected_main == "📚 LEARN":
    st.sidebar.markdown("<div style='margin-top:8px; margin-bottom:4px; font-weight:800; font-size:0.85rem; color:#2EC4B6; letter-spacing:0.05em;'>SELECT SYLLABUS TOPIC</div>", unsafe_allow_html=True)

    with st.sidebar.expander("📚 Unit I: Numbers & Complex", expanded=True):
        if st.button("🔢 Integers & Divisibility", use_container_width=True, key="nav_u1_div"):
            st.session_state.selected_topic_key = "divisibility"
        if st.button("🧮 Computation of GCD (Euclid)", use_container_width=True, key="nav_u1_gcd"):
            st.session_state.selected_topic_key = "gcd"
        if st.button("📍 Complex Numbers & Polar", use_container_width=True, key="nav_u1_cx"):
            st.session_state.selected_topic_key = "complex"

    with st.sidebar.expander("📚 Unit II: Counting & Functions", expanded=True):
        if st.button("🔀 Permutations of Distinct", use_container_width=True, key="nav_u2_perm"):
            st.session_state.selected_topic_key = "perm"
        if st.button("🎲 Combinations of Distinct", use_container_width=True, key="nav_u2_comb"):
            st.session_state.selected_topic_key = "comb"
        if st.button("🔗 Functions (Injective/Surjective)", use_container_width=True, key="nav_u2_fn"):
            st.session_state.selected_topic_key = "functions"
        if st.button("🔄 Inverse Images of Sets", use_container_width=True, key="nav_u2_inv"):
            st.session_state.selected_topic_key = "inverse_image"

st.sidebar.markdown("---")
streak = st.session_state.get("streak", 0)
xp = st.session_state.get("xp", 0)
solved_cnt = st.session_state.get("problems_solved", 0)
tier_badge = "🌱 Novice" if xp < 50 else "🥉 Apprentice" if xp < 150 else "🥈 Scholar" if xp < 300 else "👑 Math Wizard"

st.sidebar.metric("🔥 Activity Streak", f"{streak} Days")
st.sidebar.metric("🧮 Problems Solved", f"{solved_cnt}")
st.sidebar.metric("⭐ Total XP", f"{xp} XP")
st.sidebar.caption(f"Rank Tier: {tier_badge}")

if st.session_state.quiz_score["total"] > 0:
    pct = round(100 * st.session_state.quiz_score["correct"] / st.session_state.quiz_score["total"])
    st.sidebar.metric("🎯 Quiz Accuracy", f"{pct}%")


# ============================================================
# PAGE: HOME (PROMINENT HERO & VISUAL TOPIC CARDS)
# ============================================================
if selected_main == "🏠 Home":
    st.markdown("""
    <div class="hero-container">
        <div class="hero-symbol-banner">∫   Σ   √   π   ∞</div>
        <div class="hero-title">🧮 MATHMATE</div>
        <div style="font-size:1.2rem; font-weight:700; color:#2EC4B6; margin-bottom:4px;">Interactive Mathematics Lab</div>
        <div class="hero-subtitle">Learn mathematics. Solve problems. Understand every step.</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="glass-card" style="padding: 24px; border: 1px solid rgba(46, 196, 182, 0.35); box-shadow: 0 10px 30px rgba(46,196,182,0.15);">
        <div style="font-weight:800; font-size:1.3rem; color:#FFFFFF; margin-bottom:4px; display:flex; align-items:center; gap:8px;">
            ✨ What would you like to solve?
        </div>
        <div style="font-size:0.88rem; color:rgba(247,245,239,0.7); margin-bottom:14px;">
            Type or paste any syllabus problem below:
        </div>
    """, unsafe_allow_html=True)

    home_q_input = st.text_area("Question Input", key="home_question_input",
                                placeholder="e.g. Find GCD of 1071 and 462 using Euclid's algorithm  OR  Convert z = 3 + 4i to polar form",
                                height=85, label_visibility="collapsed")

    st.caption("Quick sample questions:")
    sample_qs = {
        "divisibility": "Prime factorization and divisors of 360",
        "gcd": "Find GCD of 1071 and 462 using Euclid's division algorithm",
        "complex": "Convert z = 3 + 4i to polar form",
        "perm": "Permutations P(7, 3) of 7 distinct objects",
        "comb": "Choose 4 members from a group of 9 available employees",
        "functions": "Domain: 1, 2, 3. Codomain: a, b, c. Mapping: f(1)=a, f(2)=b, f(3)=a",
        "inverse_image": "Find inverse image f⁻¹({a, c}) for f: {1,2,3,4} -> {a,b,c}"
    }

    c1, c2, c3 = st.columns(3)
    c1.button("💡 GCD of 1071 and 462", use_container_width=True, on_click=set_nav, args=("📚 LEARN", "gcd", sample_qs["gcd"]))
    c2.button("💡 Complex z = 3 + 4i", use_container_width=True, on_click=set_nav, args=("📚 LEARN", "complex", sample_qs["complex"]))
    c3.button("💡 Divisors of 360", use_container_width=True, on_click=set_nav, args=("📚 LEARN", "divisibility", sample_qs["divisibility"]))

    def handle_home_solve():
        q = st.session_state.get("home_question_input", "").strip()
        if q:
            detected, _ = parse_question(q)
            st.session_state.preset_question = q
            st.session_state.selected_topic_key = detected
            st.session_state.main_nav = "📚 LEARN"

    st.markdown("<div style='margin-top: 14px;'></div>", unsafe_allow_html=True)
    st.button("🧮 Solve Problem Step-by-Step", type="primary", use_container_width=True, on_click=handle_home_solve)
    st.markdown("</div>", unsafe_allow_html=True)

    # Identity Stats Badge Row
    st.markdown("""
    <div style="display: flex; justify-content: center; gap: 2rem; margin-top: 18px; margin-bottom: 24px; flex-wrap: wrap;">
        <div style="background: rgba(255,255,255,0.05); padding: 8px 18px; border-radius: 20px; font-size: 0.88rem; font-weight: 700; color: #2EC4B6; border: 1px solid rgba(46,196,182,0.25);">
            7 Topics
        </div>
        <div style="background: rgba(255,255,255,0.05); padding: 8px 18px; border-radius: 20px; font-size: 0.88rem; font-weight: 700; color: #FFB627; border: 1px solid rgba(255,182,39,0.25);">
            ∞ Problems
        </div>
        <div style="background: rgba(255,255,255,0.05); padding: 8px 18px; border-radius: 20px; font-size: 0.88rem; font-weight: 700; color: #E63946; border: 1px solid rgba(230,57,70,0.25);">
            🧠 Quiz Arena
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 📚 UNIT I: Integers, Real Numbers & Complex Numbers")
    u1_cols = st.columns(3)

    unit1_topics = [
        ("divisibility", "🔢 Integers & Divisibility", "Prime factorization, primality testing, complete divisor lists, τ(n) & σ(n).", "3 concepts · 4 operations"),
        ("gcd", "🧮 Euclidean Algorithm & GCD", "Find GCD using Euclid's division algorithm and Extended Bézout identity.", "3 concepts · 4 operations"),
        ("complex", "📍 Complex Numbers & Polar Form", "Rectangular to polar conversion, modulus r, argument θ, and Argand plane.", "4 concepts · Interactive Graph")
    ]

    for i, (key, title, desc, meta) in enumerate(unit1_topics):
        with u1_cols[i % 3]:
            st.markdown(f"""
            <div class="topic-card">
                <div>
                    <div class="topic-icon-lg">{TOPIC_ICONS[key]}</div>
                    <span class="unit-badge unit-badge-1">UNIT I</span>
                    <div style="font-weight:800; font-size:1.1rem; color:#FFFFFF; margin-bottom:6px;">{title}</div>
                    <div style="font-size:0.86rem; color:rgba(247,245,239,0.75); line-height:1.4;">{desc}</div>
                </div>
                <div>
                    <div class="topic-meta">{meta}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.button(f"Explore →", key=f"btn_u1_{key}", use_container_width=True, on_click=set_nav, args=("📚 LEARN", key, sample_qs[key]))

    st.markdown("<div style='margin-top:24px;'></div>", unsafe_allow_html=True)
    st.markdown("### 📚 UNIT II: Basic Counting & Basics of Functions")
    u2_cols = st.columns(4)

    unit2_topics = [
        ("perm", "🔀 Permutations", "Ordered arrangements nPr, factorial products n!, and circular table arrangements.", "3 concepts · Factorial solver"),
        ("comb", "🎲 Combinations", "Unordered selections nCr, binomial coefficient properties, and group selections.", "3 concepts · Pascal symmetry"),
        ("functions", "🔗 Function Classification", "Classify domain-to-codomain mappings with interactive bipartite graph.", "3 concepts · Bipartite graph"),
        ("inverse_image", "🔄 Inverse Images of Sets", "Pre-image computation f⁻¹(S) = { x ∈ A | f(x) ∈ S } for subsets S ⊆ B.", "2 concepts · Subset mapping")
    ]

    for i, (key, title, desc, meta) in enumerate(unit2_topics):
        with u2_cols[i % 4]:
            st.markdown(f"""
            <div class="topic-card">
                <div>
                    <div class="topic-icon-lg">{TOPIC_ICONS[key]}</div>
                    <span class="unit-badge unit-badge-2">UNIT II</span>
                    <div style="font-weight:800; font-size:1.05rem; color:#FFFFFF; margin-bottom:6px;">{title}</div>
                    <div style="font-size:0.85rem; color:rgba(247,245,239,0.75); line-height:1.4;">{desc}</div>
                </div>
                <div>
                    <div class="topic-meta">{meta}</div>
                </div>
            </div>
            """, unsafe_allow_html=True)
            st.button(f"Explore →", key=f"btn_u2_{key}", use_container_width=True, on_click=set_nav, args=("📚 LEARN", key, sample_qs[key]))


# ============================================================
# PAGE: SOLVER ENGINE (PROGRESSIVE DISCLOSURE REDESIGN)
# ============================================================
elif selected_main == "📚 LEARN":
    topic_key = st.session_state.get("selected_topic_key", "gcd")
    unit_label = TOPIC_UNITS.get(topic_key, "UNIT I")
    topic_title = TOPICS.get(topic_key, "Solver")

    st.markdown(f'<span class="unit-badge unit-badge-1">{unit_label}</span>', unsafe_allow_html=True)
    st.markdown(f'<div class="hero-title" style="font-size: 2.2rem;">{topic_title}</div>', unsafe_allow_html=True)

    initial_q = st.session_state.pop("preset_question", "")
    question_text = st.text_area("Enter or edit your problem statement", value=initial_q,
                                  placeholder=f"Type your problem for {topic_title}...",
                                  height=75)

    if question_text:
        detected_topic, parsed_params = parse_question(question_text)
    else:
        detected_topic, parsed_params = topic_key, {}

    if question_text and detected_topic != topic_key:
        detected_title = TOPICS.get(detected_topic, detected_topic)
        st.warning(f"🔍 Question detected for **{detected_title}**!")
        if st.button(f"🚀 Switch to {detected_title} & Solve Now", type="primary"):
            st.session_state.preset_question = question_text
            st.session_state.selected_topic_key = detected_topic
            st.rerun()

    steps, answer, extra = None, None, {}
    auto_trigger = bool(question_text)
    domain_str, codomain_str, mapping = "1,2,3,4", "a,b,c", {}
    target_set_str = "a,c"

    # Input controls based on topic
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    if topic_key == "divisibility":
        default_n = parsed_params.get("n", 360) if topic_key == detected_topic else 360
        n_val = st.number_input("Integer n", value=int(default_n), step=1, min_value=2)
        if st.button("🧮 Solve Problem", type="primary", use_container_width=True) or auto_trigger:
            steps, answer, extra = solve_divisibility(n_val)
            if not question_text:
                question_text = f"Find prime factorization, complete divisors, τ({n_val}), and sum of divisors σ({n_val}) for integer n = {n_val}."

    elif topic_key == "gcd":
        is_wp = parsed_params.get("is_word_problem", False) if topic_key == detected_topic else False
        nums = parsed_params.get("nums", []) if topic_key == detected_topic else []
        if (question_text and is_wp) or len(nums) > 2:
            if st.button("🧮 Solve Word Problem", type="primary", use_container_width=True) or auto_trigger:
                steps, answer, extra = solve_gcd_word_problem(question_text)
        else:
            c1, c2 = st.columns(2)
            default_a = parsed_params.get("a", 1071) if topic_key == detected_topic else 1071
            default_b = parsed_params.get("b", 462) if topic_key == detected_topic else 462
            a_val = c1.number_input("Integer a", value=int(default_a), step=1)
            b_val = c2.number_input("Integer b", value=int(default_b), step=1)
            if st.button("🧮 Solve Problem", type="primary", use_container_width=True) or auto_trigger:
                if len(nums) > 2:
                    steps, answer, extra = solve_gcd_multi(nums)
                else:
                    steps, answer, extra = solve_gcd_euclidean(a_val, b_val)
                if not question_text:
                    question_text = f"Compute gcd({a_val}, {b_val}) using Euclid's division algorithm and Extended Bézout identity."

    elif topic_key == "complex":
        c1, c2 = st.columns(2)
        default_a = parsed_params.get("a", 3.0) if topic_key == detected_topic else 3.0
        default_b = parsed_params.get("b", 4.0) if topic_key == detected_topic else 4.0
        a_val = c1.number_input("Real part (a)", value=float(default_a))
        b_val = c2.number_input("Imaginary part (b)", value=float(default_b))
        if st.button("📍 Convert to Polar Form", type="primary", use_container_width=True) or auto_trigger:
            steps, answer, extra = solve_complex_to_polar(a_val, b_val)
            if not question_text:
                question_text = f"Convert z = {a_val} + {b_val}i to polar form r(cos θ + i sin θ) and calculate modulus r and argument θ."

    elif topic_key == "perm":
        c1, c2 = st.columns(2)
        default_n = parsed_params.get("n", 7) if topic_key == detected_topic else 7
        default_r = parsed_params.get("r", 3) if topic_key == detected_topic else 3
        n_val = c1.number_input("n (total distinct objects)", value=int(default_n), step=1, min_value=0)
        r_val = c2.number_input("r (arranged objects)", value=int(default_r), step=1, min_value=0)
        if st.button("🔀 Calculate Permutations P(n,r)", type="primary", use_container_width=True) or auto_trigger:
            steps, answer, extra = solve_perm(int(n_val), int(r_val))
            if not question_text:
                question_text = f"Find number of permutations P({int(n_val)}, {int(r_val)}) for arranging {int(r_val)} objects from {int(n_val)} distinct objects."

    elif topic_key == "comb":
        c1, c2 = st.columns(2)
        default_n = parsed_params.get("n", 9) if topic_key == detected_topic else 9
        default_r = parsed_params.get("r", 4) if topic_key == detected_topic else 4
        n_val = c1.number_input("n (total items)", value=int(default_n), step=1, min_value=0)
        r_val = c2.number_input("r (chosen items)", value=int(default_r), step=1, min_value=0)
        if st.button("🎲 Calculate Combinations C(n,r)", type="primary", use_container_width=True) or auto_trigger:
            steps, answer, extra = solve_comb(int(n_val), int(r_val))
            if not question_text:
                question_text = f"Find number of combinations C({int(n_val)}, {int(r_val)}) for choosing {int(r_val)} items from {int(n_val)} distinct items."

    elif topic_key == "functions":
        st.markdown("**Domain A & Codomain B Setup:**")
        c1, c2 = st.columns(2)
        domain_str = c1.text_input("Domain A elements (comma-separated)", "1,2,3")
        codomain_str = c2.text_input("Codomain B elements (comma-separated)", "a,b,c")
        domain = [x.strip() for x in domain_str.split(",") if x.strip()]
        codomain = [x.strip() for x in codomain_str.split(",") if x.strip()]

        st.markdown("**Define Function Mapping:**")
        mapping = {}
        mcols = st.columns(min(len(domain), 4) or 1)
        for i, d in enumerate(domain):
            with mcols[i % len(mcols)]:
                mapping[d] = st.selectbox(f"f({d}) →", codomain, key=f"fn_map_{d}")

        if st.button("🔗 Analyze Function Properties", type="primary", use_container_width=True) or (auto_trigger and len(domain) > 0):
            steps, answer, extra = solve_functions(domain, codomain, mapping)
            if not question_text:
                map_str = ", ".join(f"f({k})={v}" for k, v in mapping.items())
                question_text = f"Given f: {{{', '.join(domain)}}} → {{{', '.join(codomain)}}} with mapping {map_str}, classify if f is Injective, Surjective, or Bijective."

    elif topic_key == "inverse_image":
        c1, c2, c3 = st.columns(3)
        domain_str = c1.text_input("Domain A elements", "1,2,3,4")
        codomain_str = c2.text_input("Codomain B elements", "a,b,c")
        target_set_str = c3.text_input("Target Subset S ⊆ B", "a,c")
        domain = [x.strip() for x in domain_str.split(",") if x.strip()]
        codomain = [x.strip() for x in codomain_str.split(",") if x.strip()]
        target_set = [x.strip() for x in target_set_str.split(",") if x.strip()]

        st.markdown("**Define Function Mapping:**")
        mapping = {}
        mcols = st.columns(min(len(domain), 4) or 1)
        for i, d in enumerate(domain):
            with mcols[i % len(mcols)]:
                mapping[d] = st.selectbox(f"f({d}) →", codomain, key=f"inv_map_{d}")

        if st.button("🔄 Compute Inverse Image f⁻¹(S)", type="primary", use_container_width=True) or (auto_trigger and len(domain) > 0):
            steps, answer, extra = solve_inverse_image(domain, codomain, mapping, target_set)
            if not question_text:
                map_str = ", ".join(f"f({k})={v}" for k, v in mapping.items())
                question_text = f"Given f: {{{', '.join(domain)}}} → {{{', '.join(codomain)}}} with mapping {map_str}, find the inverse image f⁻¹({{{', '.join(target_set)}}})."

    st.markdown('</div>', unsafe_allow_html=True)

    # Cache solution state
    if steps:
        st.session_state.active_solution = {
            "topic_key": topic_key,
            "question_text": question_text,
            "steps": steps,
            "answer": answer,
            "extra": extra,
            "domain_str": domain_str,
            "codomain_str": codomain_str,
            "mapping": mapping,
        }
    elif "active_solution" in st.session_state:
        act = st.session_state.active_solution
        if act.get("topic_key") == topic_key:
            question_text = act.get("question_text", question_text)
            steps = act.get("steps")
            answer = act.get("answer")
            extra = act.get("extra", {})
            domain_str = act.get("domain_str", domain_str)
            codomain_str = act.get("codomain_str", codomain_str)
            mapping = act.get("mapping", mapping)

    # RENDER REDESIGNED SOLUTION SECTION WITH PROGRESSIVE DISCLOSURE
    if steps:
        record_user_activity(solved_problem=True)
        st.markdown("---")

        # Top Section: 2 Columns for Solution Steps & Final Answer Card
        sol_col, ans_col = st.columns([7, 5])

        with sol_col:
            st.markdown("### 📚 STEP-BY-STEP SOLUTION")
            render_step_timeline(steps)

        with ans_col:
            render_answer_card(answer, latex_ans=extra.get("latex_ans"), question_text=question_text, topic_name=TOPICS.get(topic_key, topic_key), steps=steps)

            # Special status pills for Functions classification
            if topic_key == "functions":
                inj = extra.get("is_injective")
                surj = extra.get("is_surjective")
                bij = extra.get("is_bijective")

                st.markdown(f"""
                <div class="glass-card" style="padding: 16px; margin-top: 10px;">
                    <div style="font-weight: 700; font-size: 0.95rem; margin-bottom: 8px; color: #FFFFFF;">FUNCTION TYPE SUMMARY</div>
                    <div>
                        <span class="status-pill {'status-success' if inj else 'status-danger'}">{'✓ Injective' if inj else '❌ Not Injective'}</span>
                        <span class="status-pill {'status-success' if surj else 'status-danger'}">{'✓ Surjective' if surj else '❌ Not Surjective'}</span>
                        <span class="status-pill {'status-success' if bij else 'status-danger'}">{'✓ Bijective' if bij else '❌ Not Bijective'}</span>
                    </div>
                </div>
                """, unsafe_allow_html=True)

            # Complex Numbers Flow Cards
            elif topic_key == "complex":
                a = extra.get("a", 0)
                b = extra.get("b", 0)
                r = extra.get("r", 0)
                deg = extra.get("theta_deg", 0)

                st.markdown(f"""
                <div class="glass-card" style="padding: 16px; text-align: center; margin-top: 10px;">
                    <div style="font-size:0.8rem; color:#2EC4B6; font-weight:700;">CONVERSION FLOW</div>
                    <div style="font-weight:800; font-size:1.1rem; color:#FFFFFF; margin-top:4px;">z = {a} + {b}i</div>
                    <div style="color:#FFB627; font-size:0.9rem; margin:4px 0;">↓ Modulus & Argument</div>
                    <div style="font-weight:700; color:#2EC4B6;">r = {r:.2f} · θ = {deg:.2f}°</div>
                    <div style="color:#FFB627; font-size:0.9rem; margin:4px 0;">↓ Polar Form</div>
                    <div style="font-weight:800; color:#FFFFFF;">z = {r:.2f}(cos {deg:.2f}° + i sin {deg:.2f}°)</div>
                </div>
                """, unsafe_allow_html=True)

            # Save solution record in DB
            last_solved_key = f"{question_text}_{topic_key}_{answer}"
            if st.session_state.get("last_solved_key") != last_solved_key:
                st.session_state.last_solved_key = last_solved_key
                st.session_state.xp += 15
                db.save_solution(question_text or f"{TOPICS[topic_key]} problem", TOPICS[topic_key], answer, steps)

        # PROGRESSIVE DISCLOSURE EXPANDERS BELOW MAIN RESULT
        st.markdown("<div style='margin-top: 20px;'></div>", unsafe_allow_html=True)

        with st.expander("📊 Interactive Visualizations & Data Tables", expanded=True):
            if topic_key == "complex":
                fig = plot_argand_diagram_plotly(extra.get("a", 1.0), extra.get("b", 1.0))
                st.plotly_chart(fig, use_container_width=True, config={'responsive': True, 'displayModeBar': False})
            elif topic_key in ["functions", "inverse_image"]:
                fig = plot_function_diagram_plotly(
                    [x.strip() for x in domain_str.split(",") if x.strip()],
                    [x.strip() for x in codomain_str.split(",") if x.strip()],
                    mapping,
                    highlight_target_set=extra.get("target_set", [])
                )
                st.plotly_chart(fig, use_container_width=True, config={'responsive': True, 'displayModeBar': False})
            elif topic_key == "gcd" and "table_rows" in extra:
                st.markdown("**📋 Extended Euclidean Bézout Table:**")
                render_euclidean_html_table(extra["table_rows"])
            elif topic_key == "divisibility":
                st.markdown(f"**Divisors of {extra.get('n')}:**")
                st.write(extra.get("divisors", []))

        with st.expander("💡 Understand the Concept & Core Theory", expanded=False):
            render_understand_panel(topic_key)

        with st.expander("📥 Export Solution & Raw Code", expanded=False):
            full_text = format_full_solution_text(question_text, TOPICS.get(topic_key, topic_key), steps, answer)
            st.code(full_text, language=None)

            c_doc, c_pdf = st.columns(2)
            with c_doc:
                if DOCX_AVAILABLE:
                    docx_buf = build_docx(question_text, TOPICS.get(topic_key, topic_key), steps, answer)
                    if docx_buf:
                        st.download_button(
                            label="📄 Download Word (.docx)",
                            data=docx_buf,
                            file_name="MathMate_solution.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                else:
                    st.caption("Install python-docx for Word export")

            with c_pdf:
                if REPORTLAB_AVAILABLE:
                    pdf_buf = build_pdf(question_text, TOPICS.get(topic_key, topic_key), steps, answer)
                    if pdf_buf:
                        st.download_button(
                            label="📥 Download PDF (.pdf)",
                            data=pdf_buf,
                            file_name="MathMate_solution.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                else:
                    st.caption("Install reportlab for PDF export")


# ============================================================
# PAGE: PRACTICE ARENA (QUIZ)
# ============================================================
elif selected_main == "🧠 Practice Arena":
    st.markdown('<div class="hero-title">🧠 PRACTICE ARENA</div>', unsafe_allow_html=True)
    st.caption("Infinite procedurally-generated math problems across Unit I & Unit II syllabus topics. Earn XP and level up!")

    if "quiz_q" not in st.session_state:
        st.session_state.quiz_q = generate_procedural_question()

    q = st.session_state.quiz_q
    q_topic_title = TOPICS.get(q["topic"], q["topic"])
    total_q = st.session_state.quiz_score["total"]

    st.markdown(f'<span class="unit-badge unit-badge-1">{q_topic_title}</span>', unsafe_allow_html=True)
    st.markdown(f"""
    <div class="glass-card">
        <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:10px;">
            <span style="font-weight:700; color:#2EC4B6; font-size:0.9rem;">{q_topic_title}</span>
            <span style="font-size:0.85rem; color:rgba(247,245,239,0.6); font-weight:600;">Question #{total_q + 1}</span>
        </div>
        <div style="font-weight:700; font-size:1.2rem; color:#FFFFFF; margin-bottom:14px;">{q['q']}</div>
    </div>
    """, unsafe_allow_html=True)

    choice = st.radio("Choose an answer", q["options"], key="quiz_choice", label_visibility="collapsed")

    c1, c2 = st.columns(2)
    if c1.button("Submit Answer", type="primary", use_container_width=True):
        st.session_state.quiz_score["total"] += 1
        record_user_activity(solved_problem=False)

        if choice == q["answer"]:
            st.session_state.quiz_score["correct"] += 1
            st.session_state.xp += 20
            st.markdown("""
            <div style="background: rgba(46, 196, 182, 0.2); border: 1px solid #2EC4B6; border-radius: 12px; padding: 14px; text-align: center; margin-bottom: 12px;">
                <div style="font-size: 1.4rem; font-weight: 800; color: #2EC4B6;">🎉 CORRECT!</div>
                <div style="font-size: 0.9rem; color: #FFFFFF;">+20 XP Gained</div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div style="background: rgba(230, 57, 70, 0.2); border: 1px solid #E63946; border-radius: 12px; padding: 14px; text-align: center; margin-bottom: 12px;">
                <div style="font-size: 1.1rem; font-weight: 800; color: #E63946;">Not quite!</div>
                <div style="font-size: 0.9rem; color: #FFFFFF;">Correct Answer is <b>{q['answer']}</b></div>
            </div>
            """, unsafe_allow_html=True)

        with st.expander("💡 View Step-by-Step Solution Breakdown", expanded=True):
            st.write(q["exp"])

    if c2.button("Next Question →", use_container_width=True):
        st.session_state.quiz_q = generate_procedural_question()
        st.rerun()

    st.markdown("---")
    correct = st.session_state.quiz_score["correct"]
    st.metric("Total Quiz Score", f"{correct} / {total_q}" if total_q else "0 / 0")


# ============================================================
# PAGE: SEARCHABLE FORMULA BANK
# ============================================================
elif selected_main == "📐 Formula Bank":
    st.markdown('<div class="hero-title">📐 FORMULA BANK</div>', unsafe_allow_html=True)
    st.caption("Searchable reference bank for all Unit I and Unit II formulas, identities, and core theorems.")

    c_search, c_filter = st.columns([3, 1])
    search_query = c_search.text_input("🔍 Search formulas...", placeholder="e.g. Euclid, Polar, Permutation, Inverse Image...", label_visibility="collapsed")
    unit_filter = c_filter.selectbox("Filter Unit", ["All", "Unit I", "Unit II"], label_visibility="collapsed")

    filtered_concepts = {}
    for key, data in TOPIC_CONCEPTS.items():
        if unit_filter != "All" and data.get("unit") != unit_filter:
            continue
        if search_query:
            q_lower = search_query.lower()
            if not (q_lower in data["title"].lower() or q_lower in data["desc"].lower() or any(q_lower in pt.lower() for pt in data["key_points"])):
                continue
        filtered_concepts[key] = data

    if not filtered_concepts:
        st.info("No matching formulas found. Try adjusting your search query.")
    else:
        for key, c in filtered_concepts.items():
            st.markdown(f"""
            <div class="glass-card">
                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:8px;">
                    <span style="font-weight:800; font-size:1.15rem; color:#FFFFFF;">{TOPIC_ICONS[key]} {c['title']}</span>
                    <span class="unit-badge unit-badge-1">{c.get('unit','Unit I')}</span>
                </div>
                <div style="font-size:0.88rem; color:rgba(247,245,239,0.75); margin-bottom:12px;">{c['desc']}</div>
            </div>
            """, unsafe_allow_html=True)

            with st.expander(f"📌 {c['title']} Formulas & Identities", expanded=True):
                st.markdown("**Main Formula:**")
                st.latex(c['formula'])
                st.markdown("**Key Identity / Relationship:**")
                st.latex(c['identity'])
                st.markdown("**Key Properties:**")
                for pt in c['key_points']:
                    st.markdown(f"- {pt}")
                st.code(c['formula'], language=None)


# ============================================================
# PAGE: HISTORY & PROGRESS DASHBOARD
# ============================================================
elif selected_main == "📜 Solution History":
    st.markdown('<div class="hero-title">📜 YOUR PROGRESS DASHBOARD</div>', unsafe_allow_html=True)
    st.caption("Review your solved problems and tracking metrics (Persisted in Database).")

    m1, m2, m3, m4 = st.columns(4)
    history = db.fetch_history(limit=50)

    m1.metric("Problems Solved", f"{st.session_state.get('problems_solved', len(history))}")
    m2.metric("Total XP", f"{st.session_state.get('xp', 0)} XP")

    q_tot = st.session_state.quiz_score["total"]
    q_pct = f"{round(100 * st.session_state.quiz_score['correct'] / q_tot)}%" if q_tot > 0 else "0%"
    m3.metric("Quiz Accuracy", q_pct)
    m4.metric("Activity Streak", f"{st.session_state.get('streak', 0)} Days 🔥")

    st.markdown("---")
    st.markdown("### 📜 RECENT SOLUTIONS")

    c_filter, c_search = st.columns([2, 2])
    topic_filter = c_filter.selectbox("Filter by Topic", ["All"] + list(TOPICS.values()))

    filtered_history = db.fetch_history(limit=50, topic_filter=topic_filter if topic_filter != "All" else None)

    if not filtered_history:
        st.info("No solved questions in history yet — head to **Home** or pick a topic in **LEARN** to get started.")
    else:
        for item in filtered_history:
            st.markdown(f"""
            <div class="glass-card">
                <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:6px;">
                    <span style="font-weight:700; color:#2EC4B6; font-size:0.95rem;">{item.get('topic','')}</span>
                    <span style="font-size:0.82rem; color:rgba(247,245,239,0.5);">{item.get('time','')}</span>
                </div>
                <div style="font-size:1.05rem; font-weight:600; color:#FFFFFF; margin-bottom:6px;">{item.get('question','')}</div>
                <div style="font-size:0.9rem; color:#FFB627; font-weight:700;">Answer: {item.get('answer','')}</div>
            </div>
            """, unsafe_allow_html=True)
            if item.get("steps"):
                with st.expander("🔍 View Step-by-Step Reasoning", expanded=False):
                    render_step_timeline(item.get("steps"))
