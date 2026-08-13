"""
MathMate — Interactive Mathematics Lab
Streamlit app covering: Euclidean Algorithm & GCD, Complex Numbers & Polar Form,
De Moivre's Theorem, Permutations & Combinations, Injective/Surjective/Bijective
Functions, and Limits & Continuity.
"""

import io
import re
import math
import random
from datetime import datetime

import streamlit as st
import numpy as np
import matplotlib.pyplot as plt
import sympy as sp
from sympy import symbols, sympify, limit, oo, latex, I, re as s_re, im as s_im, Rational
from sympy.parsing.sympy_parser import (
    parse_expr, standard_transformations, implicit_multiplication_application,
)

# Optional OCR — degrade gracefully if not installed / no tesseract binary on host
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

# Optional export libs — degrade gracefully
try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_AVAILABLE = True
except Exception:
    REPORTLAB_AVAILABLE = False


# ============================================================
# PAGE CONFIG & THEME
# ============================================================
st.set_page_config(page_title="MathMate — Interactive Mathematics Lab",
                    page_icon="🧮", layout="wide")

INK = "#14213D"
PAPER = "#F7F5EF"
TEAL = "#2EC4B6"
AMBER = "#FFB627"
CORAL = "#E63946"
GRAPHITE = "#4A4E69"

st.markdown(f"""
<style>
.stApp {{ background-color: {INK}; }}
h1, h2, h3, h4, p, span, label, .stMarkdown {{ color: {PAPER}; }}
.step-box {{
    background: rgba(247,245,239,0.06); border: 1px solid rgba(247,245,239,0.15);
    border-left: 4px solid {TEAL}; border-radius: 8px; padding: 14px 18px; margin: 10px 0;
}}
.step-title {{ color: {TEAL}; font-weight: 700; font-size: 0.85rem; letter-spacing: .03em;
    text-transform: uppercase; margin-bottom: 4px; }}
.result-box {{
    background: rgba(46,196,182,0.12); border: 1px solid {TEAL}; border-radius: 10px;
    padding: 16px 20px; font-size: 1.15rem; font-weight: 600; color: {TEAL};
}}
.topic-badge {{
    display: inline-block; background: rgba(255,182,39,0.15); color: {AMBER};
    border: 1px solid rgba(255,182,39,0.4); border-radius: 999px; padding: 4px 14px;
    font-size: 0.8rem; font-weight: 700; letter-spacing: .04em;
}}
[data-testid="stSidebar"] {{ background-color: #0D1730; }}
</style>
""", unsafe_allow_html=True)


# ============================================================
# TOPIC DEFINITIONS + DETECTION
# ============================================================
TOPICS = {
    "gcd": "🧮 Euclidean Algorithm & GCD",
    "complex": "📍 Complex Numbers & Polar Form",
    "demoivre": "🔄 De Moivre's Theorem",
    "permcomb": "🔢 Permutations & Combinations",
    "functions": "🔗 Injective, Surjective & Bijective Functions",
    "limits": "📈 Limits & Continuity",
}

TOPIC_KEYWORDS = {
    "gcd": ["gcd", "hcf", "euclidean", "greatest common divisor", "highest common factor"],
    "complex": ["polar form", "modulus", "argument", "rectangular form", "complex number", "argand"],
    "demoivre": ["de moivre", "demoivre", "nth root", "z^n", "power of complex", "roots of unity"],
    "permcomb": ["permutation", "combination", "arrange", "select", "ncr", "npr", "how many ways"],
    "functions": ["injective", "surjective", "bijective", "one-one", "onto", "into", "mapping"],
    "limits": ["limit", "lim ", "continuity", "continuous", "discontinuous", "x->", "x→"],
}


def detect_topic(text: str) -> str:
    """Very lightweight keyword-based topic classifier with a scoring fallback."""
    t = text.lower()
    scores = {k: 0 for k in TOPICS}
    for topic, kws in TOPIC_KEYWORDS.items():
        for kw in kws:
            if kw in t:
                scores[topic] += 1
    # heuristics for symbol-only inputs
    if re.search(r"\bi\b|√-?1|imaginary", t):
        scores["complex"] += 1
    if re.search(r"\(.*\)\^\d|\^n\b", t):
        scores["demoivre"] += 1
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else None


# ============================================================
# SOLVERS — each returns (steps: list[(title, body)], final_answer: str, extra: dict)
# ============================================================

def solve_gcd(a: int, b: int):
    steps = []
    a, b = abs(int(a)), abs(int(b))
    x, y = a, b
    if y == 0:
        return [("Edge case", f"gcd({x}, 0) = {x}")], str(x), {}
    steps.append(("Set up", f"Apply the Euclidean Algorithm to gcd({x}, {y})."))
    while y != 0:
        q = x // y
        r = x % y
        steps.append((f"Divide", f"{x} = {q} × {y} + {r}"))
        x, y = y, r
    steps.append(("Remainder is 0", f"Stop — the last non-zero remainder is the GCD."))
    return steps, str(x), {"pairs": (a, b)}


def solve_complex_to_polar(a: float, b: float):
    steps = []
    steps.append(("Identify the form", f"z = {a} + {b}i is in rectangular (x + yi) form."))
    r = math.sqrt(a**2 + b**2)
    steps.append(("Find the modulus", f"r = √(x² + y²) = √({a}² + {b}²) = √{a**2 + b**2:g} = {r:.4f}"))
    theta = math.atan2(b, a)
    theta_deg = math.degrees(theta)
    quadrant = "1st" if a > 0 and b >= 0 else "2nd" if a < 0 and b >= 0 else "3rd" if a < 0 and b < 0 else "4th"
    steps.append(("Find the argument", f"θ = atan2(y, x) = {theta:.4f} rad ≈ {theta_deg:.2f}°  ({quadrant} quadrant)"))
    steps.append(("Write in polar form", f"z = r(cosθ + i sinθ) = {r:.4f}(cos({theta_deg:.2f}°) + i·sin({theta_deg:.2f}°))"))
    answer = f"z = {r:.4f} · (cos {theta_deg:.2f}° + i sin {theta_deg:.2f}°)"
    return steps, answer, {"r": r, "theta": theta, "a": a, "b": b}


def solve_demoivre(a: float, b: float, n: int):
    steps = []
    steps.append(("Convert to polar", f"Start from z = {a} + {b}i and convert to polar form first."))
    r = math.sqrt(a**2 + b**2)
    theta = math.atan2(b, a)
    steps.append(("Modulus & argument", f"r = {r:.4f},  θ = {math.degrees(theta):.2f}°"))
    steps.append(("Apply De Moivre's theorem", f"zⁿ = rⁿ(cos(nθ) + i sin(nθ)),  n = {n}"))
    r_n = r ** n
    theta_n = theta * n
    theta_n_deg = math.degrees(theta_n) % 360
    steps.append(("Compute rⁿ and nθ", f"rⁿ = {r:.4f}^{n} = {r_n:.4f}\nnθ = {n} × {math.degrees(theta):.2f}° = {theta_n_deg:.2f}° (mod 360°)"))
    real_part = r_n * math.cos(theta_n)
    imag_part = r_n * math.sin(theta_n)
    steps.append(("Convert back to rectangular", f"zⁿ = {r_n:.4f}(cos {theta_n_deg:.2f}° + i sin {theta_n_deg:.2f}°) = {real_part:.4f} + {imag_part:.4f}i"))
    answer = f"z^{n} ≈ {real_part:.4f} + {imag_part:.4f}i"
    return steps, answer, {"r_n": r_n, "theta_n": theta_n, "real": real_part, "imag": imag_part}


def demoivre_roots(a: float, b: float, n: int):
    """All n-th roots of z = a+bi."""
    r = math.sqrt(a**2 + b**2)
    theta = math.atan2(b, a)
    r_root = r ** (1 / n)
    roots = []
    for k in range(n):
        angle = (theta + 2 * math.pi * k) / n
        roots.append((r_root * math.cos(angle), r_root * math.sin(angle), math.degrees(angle)))
    return roots, r_root


def solve_permcomb(kind: str, n: int, r: int):
    steps = []
    n, r = int(n), int(r)
    if r > n:
        return [("Invalid input", "r cannot be greater than n.")], "Undefined", {}
    if kind == "Permutation (nPr)":
        steps.append(("Formula", "nPr = n! / (n − r)!"))
        steps.append(("Substitute", f"{n}P{r} = {n}! / ({n} − {r})! = {n}! / {n-r}!"))
        val = math.perm(n, r)
        steps.append(("Simplify", f"= {' × '.join(str(i) for i in range(n, n-r, -1))} = {val}"))
        answer = str(val)
    else:
        steps.append(("Formula", "nCr = n! / (r! (n − r)!)"))
        steps.append(("Substitute", f"{n}C{r} = {n}! / ({r}! × {n-r}!)"))
        val = math.comb(n, r)
        steps.append(("Simplify", f"= {val}"))
        answer = str(val)
    return steps, answer, {}


def solve_functions(domain: list, codomain: list, mapping: dict):
    """mapping: dict domain_element -> codomain_element"""
    steps = []
    steps.append(("Set up", f"Domain = {domain}, Codomain = {codomain}, mapping f = {mapping}"))
    images = list(mapping.values())
    is_injective = len(images) == len(set(images))
    steps.append(("Check injectivity (one-one)",
                  "No two distinct domain elements share an image." if is_injective
                  else "Two or more domain elements map to the same image → not injective."))
    is_surjective = set(codomain) == set(images)
    unmapped = set(codomain) - set(images)
    steps.append(("Check surjectivity (onto)",
                  "Every element of the codomain is hit by some domain element." if is_surjective
                  else f"Codomain element(s) {sorted(unmapped)} have no pre-image → not surjective."))
    is_bijective = is_injective and is_surjective
    steps.append(("Conclusion", "Bijective (both injective and surjective)." if is_bijective
                  else "Not bijective."))
    classification = []
    if is_injective: classification.append("Injective")
    if is_surjective: classification.append("Surjective")
    if is_bijective: classification.append("Bijective")
    answer = ", ".join(classification) if classification else "Neither injective nor surjective"
    return steps, answer, {"injective": is_injective, "surjective": is_surjective, "bijective": is_bijective}


def solve_limit(expr_str: str, var_str: str, point_str: str):
    steps = []
    x = symbols(var_str)
    transformations = standard_transformations + (implicit_multiplication_application,)
    expr = parse_expr(expr_str.replace("^", "**"), transformations=transformations)
    point = oo if point_str.strip() in ("inf", "infinity", "oo") else \
            -oo if point_str.strip() in ("-inf", "-infinity", "-oo") else sympify(point_str)

    steps.append(("Original expression", f"lim_{{{var_str}→{point_str}}} {expr}"))

    # try direct substitution first
    try:
        direct = expr.subs(x, point)
        if direct.is_finite and not direct.has(sp.zoo, sp.nan):
            steps.append(("Direct substitution", f"Substitute {var_str} = {point_str}: result is finite → {direct}"))
            result = direct
        else:
            raise ValueError("indeterminate")
    except Exception:
        steps.append(("Direct substitution", f"Substituting {var_str} = {point_str} gives an indeterminate form (0/0 or ∞/∞) — simplify or apply L'Hôpital's rule."))
        result = limit(expr, x, point)
        simplified = sp.simplify(expr)
        if simplified != expr:
            steps.append(("Simplify", f"Simplify the expression: {simplified}"))
        steps.append(("Evaluate the limit", f"L = {result}"))

    # continuity check at the point (only for finite points)
    continuity_note = None
    if point not in (oo, -oo):
        try:
            f_val = expr.subs(x, point)
            if sp.simplify(f_val - result) == 0 and f_val.is_finite:
                continuity_note = f"f({var_str}) is continuous at {var_str} = {point_str} since lim = f({point_str}) = {result}."
            else:
                continuity_note = f"f({var_str}) is NOT continuous at {var_str} = {point_str} (limit ≠ function value, or function undefined there)."
        except Exception:
            continuity_note = None
    steps.append(("Final result", f"lim_{{{var_str}→{point_str}}} {expr} = {result}"))
    answer = str(result)
    return steps, answer, {"continuity_note": continuity_note, "expr": expr, "var": x, "point": point}


# ============================================================
# VISUALIZATIONS
# ============================================================

def plot_complex_plane(points: list, labels: list, colors: list = None):
    fig, ax = plt.subplots(figsize=(4, 4))
    fig.patch.set_facecolor(INK)
    ax.set_facecolor(INK)
    colors = colors or [TEAL, AMBER, CORAL, "#9B5DE5"]
    max_r = max([abs(complex(*p)) for p in points] + [1]) * 1.3
    ax.axhline(0, color=PAPER, alpha=0.3, linewidth=1)
    ax.axvline(0, color=PAPER, alpha=0.3, linewidth=1)
    for i, (p, lab) in enumerate(zip(points, labels)):
        c = colors[i % len(colors)]
        ax.plot([0, p[0]], [0, p[1]], color=c, linewidth=1.5)
        ax.scatter([p[0]], [p[1]], color=c, s=40, zorder=5)
        ax.annotate(lab, (p[0], p[1]), textcoords="offset points", xytext=(6, 6), color=c, fontsize=10)
    ax.set_xlim(-max_r, max_r)
    ax.set_ylim(-max_r, max_r)
    ax.set_xlabel("Re", color=PAPER)
    ax.set_ylabel("Im", color=PAPER)
    ax.tick_params(colors=PAPER)
    for spine in ax.spines.values():
        spine.set_color(GRAPHITE)
    ax.set_aspect("equal")
    return fig


def plot_function_diagram(domain, codomain, mapping):
    fig, ax = plt.subplots(figsize=(5, 3.2))
    fig.patch.set_facecolor(INK)
    ax.set_facecolor(INK)
    ax.axis("off")
    y_dom = np.linspace(0.85, 0.15, len(domain)) if len(domain) > 1 else [0.5]
    y_cod = np.linspace(0.85, 0.15, len(codomain)) if len(codomain) > 1 else [0.5]
    dom_pos = {d: (0.15, y) for d, y in zip(domain, y_dom)}
    cod_pos = {c: (0.75, y) for c, y in zip(codomain, y_cod)}
    for d, (x, y) in dom_pos.items():
        ax.scatter([x], [y], color=TEAL, s=200, zorder=5)
        ax.annotate(str(d), (x, y), color=INK, ha="center", va="center", fontweight="bold", fontsize=9)
    for c, (x, y) in cod_pos.items():
        ax.scatter([x], [y], color=AMBER, s=200, zorder=5)
        ax.annotate(str(c), (x, y), color=INK, ha="center", va="center", fontweight="bold", fontsize=9)
    for d, c in mapping.items():
        x1, y1 = dom_pos[d]
        x2, y2 = cod_pos[c]
        ax.annotate("", xy=(x2 - 0.04, y2), xytext=(x1 + 0.04, y1),
                    arrowprops=dict(arrowstyle="->", color=PAPER, alpha=0.6))
    ax.text(0.15, 0.98, "Domain", color=PAPER, ha="center", fontsize=10)
    ax.text(0.75, 0.98, "Codomain", color=PAPER, ha="center", fontsize=10)
    return fig


def plot_limit_function(expr, var, point):
    fig, ax = plt.subplots(figsize=(5, 3.2))
    fig.patch.set_facecolor(INK)
    ax.set_facecolor(INK)
    try:
        pt = float(point) if point not in (oo, -oo) else 0
        xs = np.linspace(pt - 5, pt + 5, 400)
        f = sp.lambdify(var, expr, "numpy")
        with np.errstate(all="ignore"):
            ys = f(xs)
        ax.plot(xs, ys, color=TEAL, linewidth=2)
        ax.axvline(pt, color=AMBER, linestyle="--", alpha=0.7, label=f"{var} = {point}")
        ax.legend(facecolor=INK, labelcolor=PAPER, edgecolor=GRAPHITE)
    except Exception:
        ax.text(0.5, 0.5, "Graph unavailable for this expression", color=PAPER, ha="center", va="center")
    ax.set_facecolor(INK)
    ax.tick_params(colors=PAPER)
    for spine in ax.spines.values():
        spine.set_color(GRAPHITE)
    return fig


# ============================================================
# EXPORT HELPERS
# ============================================================

def build_docx(question, topic, steps, answer):
    doc = Document()
    doc.add_heading("MathMate — Solution", level=1)
    doc.add_paragraph(f"Topic: {topic}")
    doc.add_paragraph(f"Question: {question}")
    doc.add_heading("Step-by-step solution", level=2)
    for title, body in steps:
        doc.add_paragraph(title, style="Heading 3")
        doc.add_paragraph(body)
    doc.add_heading("Final Answer", level=2)
    doc.add_paragraph(answer)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf(question, topic, steps, answer):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph("MathMate — Solution", styles["Title"]), Spacer(1, 12)]
    story.append(Paragraph(f"<b>Topic:</b> {topic}", styles["Normal"]))
    story.append(Paragraph(f"<b>Question:</b> {question}", styles["Normal"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Step-by-step solution", styles["Heading2"]))
    for title, body in steps:
        story.append(Paragraph(f"<b>{title}</b>", styles["Normal"]))
        story.append(Paragraph(body.replace("\n", "<br/>"), styles["Normal"]))
        story.append(Spacer(1, 6))
    story.append(Paragraph("Final Answer", styles["Heading2"]))
    story.append(Paragraph(str(answer), styles["Normal"]))
    doc.build(story)
    buf.seek(0)
    return buf


def render_steps(steps):
    for title, body in steps:
        st.markdown(f"""
        <div class="step-box">
            <div class="step-title">{title}</div>
            <div style="white-space:pre-line;">{body}</div>
        </div>
        """, unsafe_allow_html=True)


def push_history(question, topic, answer):
    if "history" not in st.session_state:
        st.session_state.history = []
    st.session_state.history.insert(0, {
        "question": question, "topic": topic, "answer": answer,
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
    })


# ============================================================
# SIDEBAR NAV
# ============================================================
st.sidebar.markdown("## 🧮 MathMate")
st.sidebar.caption("Interactive Mathematics Lab")
page = st.sidebar.radio("Navigate", ["🏠 Home", "✨ Solve", "🎯 Practice & Quiz", "🕘 History"], label_visibility="collapsed")

if "streak" not in st.session_state:
    st.session_state.streak = 0
if "quiz_score" not in st.session_state:
    st.session_state.quiz_score = {"correct": 0, "total": 0}

st.sidebar.markdown("---")
st.sidebar.metric("🔥 Streak", st.session_state.streak)
if st.session_state.quiz_score["total"] > 0:
    pct = round(100 * st.session_state.quiz_score["correct"] / st.session_state.quiz_score["total"])
    st.sidebar.metric("🎯 Quiz accuracy", f"{pct}%")


# ============================================================
# PAGE: HOME
# ============================================================
if page == "🏠 Home":
    st.title("🧮 MathMate")
    st.subheader("Scan → Detect Topic → Solve Step-by-Step → Understand → Save")
    st.write("Six syllabus topics, one lab. Pick a topic below or head to **Solve** to enter a question directly.")
    cols = st.columns(3)
    for i, (key, label) in enumerate(TOPICS.items()):
        with cols[i % 3]:
            st.markdown(f"""
            <div class="step-box" style="min-height:110px;">
                <div class="step-title">{label}</div>
            </div>
            """, unsafe_allow_html=True)


# ============================================================
# PAGE: SOLVE
# ============================================================
elif page == "✨ Solve":
    st.title("Solve a question")

    input_mode = st.radio("Input method", ["Type", "Paste", "Scan (image upload)"], horizontal=True)
    question_text = ""

    if input_mode == "Scan (image upload)":
        img_file = st.file_uploader("Upload a photo of your question", type=["png", "jpg", "jpeg"])
        cam_file = st.camera_input("...or capture with your camera")
        source = img_file or cam_file
        if source:
            if OCR_AVAILABLE:
                image = Image.open(source)
                st.image(image, caption="Uploaded question", width=300)
                question_text = pytesseract.image_to_string(image)
                st.text_area("OCR result (edit if needed)", value=question_text, key="ocr_text")
                question_text = st.session_state.get("ocr_text", question_text)
            else:
                st.warning("OCR isn't available on this deployment (pytesseract/tesseract not installed). "
                           "Add `pytesseract` to requirements.txt and `tesseract-ocr` to packages.txt, "
                           "or type/paste the question below instead.")
                question_text = st.text_area("Type the question from your image", "")
    else:
        question_text = st.text_area("Enter your question",
                                      placeholder="e.g. Find z⁴ if z = 1 + i√3, expressing the answer in rectangular form.",
                                      height=100)

    detected = detect_topic(question_text) if question_text else None
    topic_options = list(TOPICS.keys())
    default_idx = topic_options.index(detected) if detected else 0
    if detected:
        st.markdown(f'<span class="topic-badge">DETECTED: {TOPICS[detected]}</span>', unsafe_allow_html=True)
    chosen = st.selectbox("Topic (auto-detected — override if needed)",
                           topic_options, index=default_idx, format_func=lambda k: TOPICS[k])
    st.markdown("---")

    steps, answer, extra = None, None, {}

    if chosen == "gcd":
        c1, c2 = st.columns(2)
        a = c1.number_input("a", value=1071, step=1)
        b = c2.number_input("b", value=462, step=1)
        if st.button("Solve step-by-step", type="primary"):
            steps, answer, extra = solve_gcd(a, b)

    elif chosen == "complex":
        c1, c2 = st.columns(2)
        a = c1.number_input("Real part (a)", value=1.0)
        b = c2.number_input("Imaginary part (b)", value=1.7320508)
        if st.button("Solve step-by-step", type="primary"):
            steps, answer, extra = solve_complex_to_polar(a, b)

    elif chosen == "demoivre":
        c1, c2, c3 = st.columns(3)
        a = c1.number_input("Real part (a)", value=1.0)
        b = c2.number_input("Imaginary part (b)", value=1.7320508)
        n = c3.number_input("Power n", value=4, step=1)
        show_roots = st.checkbox("Also show all n-th roots of z (instead of zⁿ)")
        if st.button("Solve step-by-step", type="primary"):
            if show_roots:
                roots, r_root = demoivre_roots(a, b, int(n))
                steps = [("Convert to polar", f"z = {a} + {b}i → r = {math.hypot(a,b):.4f}, θ = {math.degrees(math.atan2(b,a)):.2f}°")]
                steps.append(("Root formula", f"wₖ = r^(1/n) [cos((θ+2πk)/n) + i sin((θ+2πk)/n)],  k = 0,…,{int(n)-1}"))
                for k, (re_, im_, ang) in enumerate(roots):
                    steps.append((f"Root k={k}", f"w{k} = {re_:.4f} + {im_:.4f}i  (angle {ang:.2f}°)"))
                answer = ", ".join(f"{re_:.3f}+{im_:.3f}i" for re_, im_, _ in roots)
                extra = {"roots": roots}
            else:
                steps, answer, extra = solve_demoivre(a, b, int(n))

    elif chosen == "permcomb":
        c1, c2, c3 = st.columns(3)
        kind = c1.selectbox("Type", ["Permutation (nPr)", "Combination (nCr)"])
        n = c2.number_input("n", value=5, step=1, min_value=0)
        r = c3.number_input("r", value=2, step=1, min_value=0)
        if st.button("Solve step-by-step", type="primary"):
            steps, answer, extra = solve_permcomb(kind, n, r)

    elif chosen == "functions":
        st.caption("Define a finite function by its mapping to classify it.")
        c1, c2 = st.columns(2)
        domain_str = c1.text_input("Domain elements (comma-separated)", "1,2,3")
        codomain_str = c2.text_input("Codomain elements (comma-separated)", "a,b,c,d")
        domain = [x.strip() for x in domain_str.split(",") if x.strip()]
        codomain = [x.strip() for x in codomain_str.split(",") if x.strip()]
        st.write("Map each domain element to a codomain element:")
        mapping = {}
        mcols = st.columns(min(len(domain), 4) or 1)
        for i, d in enumerate(domain):
            with mcols[i % len(mcols)]:
                mapping[d] = st.selectbox(f"f({d}) =", codomain, key=f"map_{d}")
        if st.button("Classify function", type="primary"):
            steps, answer, extra = solve_functions(domain, codomain, mapping)

    elif chosen == "limits":
        c1, c2, c3 = st.columns(3)
        expr_str = c1.text_input("f(x) =", "sin(3*x)/x")
        var_str = c2.text_input("Variable", "x")
        point_str = c3.text_input("x →", "0")
        if st.button("Solve step-by-step", type="primary"):
            try:
                steps, answer, extra = solve_limit(expr_str, var_str, point_str)
            except Exception as e:
                st.error(f"Couldn't parse that expression: {e}")

    # ---- render result ----
    if steps:
        st.markdown("### Step-by-step solution")
        render_steps(steps)
        st.markdown(f'<div class="result-box">✅ Final answer: {answer}</div>', unsafe_allow_html=True)
        st.session_state.streak += 1
        push_history(question_text or f"{TOPICS[chosen]} problem", TOPICS[chosen], answer)

        # visualization
        st.markdown("### Visualization")
        if chosen == "complex":
            fig = plot_complex_plane([(extra["a"], extra["b"])], ["z"])
            st.pyplot(fig)
        elif chosen == "demoivre":
            if "roots" in extra:
                pts = [(re_, im_) for re_, im_, _ in extra["roots"]]
                labs = [f"w{k}" for k in range(len(pts))]
                st.pyplot(plot_complex_plane(pts, labs))
            else:
                st.pyplot(plot_complex_plane([(extra["real"], extra["imag"])], ["zⁿ"], colors=[AMBER]))
        elif chosen == "functions":
            st.pyplot(plot_function_diagram(
                [x.strip() for x in domain_str.split(",") if x.strip()],
                [x.strip() for x in codomain_str.split(",") if x.strip()],
                mapping))
        elif chosen == "limits":
            st.pyplot(plot_limit_function(extra["expr"], extra["var"], extra["point"]))
            if extra.get("continuity_note"):
                st.info(extra["continuity_note"])

        # ---- export / actions ----
        st.markdown("### Save / Export")
        e1, e2, e3, e4 = st.columns(4)
        with e1:
            st.button("📋 Copy answer", on_click=lambda: None, help="Answer shown above — select & copy.")
        with e2:
            if DOCX_AVAILABLE:
                buf = build_docx(question_text or TOPICS[chosen], TOPICS[chosen], steps, answer)
                st.download_button("⬇️ DOCX", buf, file_name="mathmate_solution.docx")
            else:
                st.caption("Add `python-docx` to enable DOCX export.")
        with e3:
            if REPORTLAB_AVAILABLE:
                buf = build_pdf(question_text or TOPICS[chosen], TOPICS[chosen], steps, answer)
                st.download_button("⬇️ PDF", buf, file_name="mathmate_solution.pdf")
            else:
                st.caption("Add `reportlab` to enable PDF export.")
        with e4:
            st.button("🔗 Share link", help="Wire this up to your own link-sharing/backend.")


# ============================================================
# PAGE: PRACTICE & QUIZ
# ============================================================
elif page == "🎯 Practice & Quiz":
    st.title("Practice & Quiz")
    QUESTION_BANK = [
        {"q": "Find gcd(48, 18) using the Euclidean algorithm.", "topic": "gcd",
         "options": ["6", "8", "12", "3"], "answer": "6"},
        {"q": "Convert z = √3 + i to polar form. What is the argument θ?", "topic": "complex",
         "options": ["30°", "45°", "60°", "90°"], "answer": "30°"},
        {"q": "If z = 2(cos30° + i sin30°), what is |z³|?", "topic": "demoivre",
         "options": ["6", "8", "4", "2"], "answer": "8"},
        {"q": "How many ways can 4 distinct books be arranged on a shelf?", "topic": "permcomb",
         "options": ["24", "12", "16", "10"], "answer": "24"},
        {"q": "A function f: {1,2,3} → {a,b} where f(1)=a, f(2)=a, f(3)=b — is it injective?", "topic": "functions",
         "options": ["No", "Yes"], "answer": "No"},
        {"q": "lim(x→0) sin(x)/x = ?", "topic": "limits",
         "options": ["1", "0", "∞", "undefined"], "answer": "1"},
    ]
    if "quiz_q" not in st.session_state:
        st.session_state.quiz_q = random.choice(QUESTION_BANK)

    q = st.session_state.quiz_q
    st.markdown(f'<span class="topic-badge">{TOPICS[q["topic"]]}</span>', unsafe_allow_html=True)
    st.markdown(f"#### {q['q']}")
    choice = st.radio("Choose an answer", q["options"], key="quiz_choice")

    c1, c2 = st.columns(2)
    if c1.button("Submit answer"):
        st.session_state.quiz_score["total"] += 1
        if choice == q["answer"]:
            st.session_state.quiz_score["correct"] += 1
            st.success("Correct! 🎉")
        else:
            st.error(f"Not quite — the correct answer is {q['answer']}.")
    if c2.button("Next question"):
        st.session_state.quiz_q = random.choice(QUESTION_BANK)
        st.rerun()

    st.markdown("---")
    total = st.session_state.quiz_score["total"]
    correct = st.session_state.quiz_score["correct"]
    st.metric("Score", f"{correct} / {total}" if total else "0 / 0")


# ============================================================
# PAGE: HISTORY
# ============================================================
elif page == "🕘 History":
    st.title("Solution history")
    history = st.session_state.get("history", [])
    if not history:
        st.info("No solved questions yet — head to **Solve** to get started.")
    else:
        for item in history:
            with st.expander(f"{item['topic']} · {item['time']}"):
                st.write(f"**Question:** {item['question']}")
                st.write(f"**Answer:** {item['answer']}")
